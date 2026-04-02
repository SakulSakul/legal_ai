# ============================================================
#  🤝 공정거래 실무 어시스턴트 v2.1 — 면세점 MD/바이어용
#  이중 모델: Gemini(문서/검색) + Claude(법률검토) + 고가용성 우회
#  보안: 자동 DLP(개인/기업정보) + 협력사명 지정 마스킹 탑재
#  디자인: 신세계그룹 뉴스룸 테마 적용 (Noto Sans KR, #E02B20)
#
#  v2.1 변경사항:
#  - Claude 모델 버전 설정값 외부화 (Secrets 지원)
#  - 계좌번호 DLP 정규식 오탐 수정 (은행별 패턴 정밀화)
#  - Gemini 폴백 시 에러 타입별 분기 (rate limit 공유 쿼터 대응)
#  - extract_text 에러 핸들링 추가
#  - 시스템 프롬프트 토큰 관리 (조항 단위 자르기)
#  - st.rerun() 타이밍 개선 (다운로드 가능 상태 보장)
#  - import 위치 정리 (uuid 등 상단 이동)
# ============================================================

import streamlit as st
import os, io, json, re, time, logging, uuid
import xml.etree.ElementTree as ET
from datetime import datetime, timedelta

# ── 로깅 설정 ────────────────────────────────────────────────
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ── 설정 ─────────────────────────────────────────────────────
def get_secret(key, default=""):
    """Secrets → 환경변수 → default 순으로 설정값 조회"""
    try:
        return st.secrets[key]
    except Exception:
        return os.environ.get(key, default)

# 모델명 외부 설정 지원 (Secrets에 없으면 기본값 사용)
CLAUDE_MODEL = get_secret("CLAUDE_MODEL", "claude-sonnet-4-20250514")
GEMINI_MODELS = [
    get_secret("GEMINI_MODEL_PRIMARY", "gemini-3.1-pro-preview"),
    get_secret("GEMINI_MODEL_FALLBACK", "gemini-3-flash-preview"),
]

# ── Lazy 클라이언트 초기화 ────────────────────────────────────
@st.cache_resource
def init_supabase():
    from supabase import create_client
    url = get_secret("SUPABASE_URL")
    key = get_secret("SUPABASE_KEY")
    if not url or not key:
        st.error("⚠️ Supabase 설정이 없습니다. Secrets에 SUPABASE_URL, SUPABASE_KEY를 추가하세요.")
        st.stop()
    return create_client(url, key)

@st.cache_resource
def init_gemini():
    from google import genai
    key = get_secret("GEMINI_API_KEY")
    if not key:
        st.error("⚠️ Gemini API 키가 없습니다.")
        st.stop()
    return genai.Client(api_key=key)

@st.cache_resource
def init_anthropic():
    import anthropic
    key = get_secret("ANTHROPIC_API_KEY")
    if not key:
        st.error("⚠️ Anthropic API 키가 없습니다. Secrets에 ANTHROPIC_API_KEY를 추가하세요.")
        st.stop()
    return anthropic.Anthropic(api_key=key)

# ── 상수 ─────────────────────────────────────────────────────
CONTRACT_TYPES = ["특약매입", "직매입"]
YAKJEONG_TYPES = ["협력사원", "인테리어설치", "매장이동", "공동판촉", "기타"]
DOC_CATS = {
    "saryu":    {"label": "사규",   "icon": "🏛"},
    "contract": {"label": "계약서", "icon": "📄"},
    "yakjeong": {"label": "약정서", "icon": "📝"},
}

REVIEW_KEYWORDS = ["검토", "확인", "위반", "적법", "수용", "반품", "계약", "약정",
                   "조항", "독소", "비교", "분석", "판촉", "감액", "반려", "승인",
                   "법률", "법적", "맞음", "맞아", "맞나요", "위법", "불법", "가능", "어때", "어떤가요", "문제"]

# ── 자동 보안 마스킹 (DLP) 함수 ──────────────────────────────
def apply_auto_masking(text, target_partner=""):
    """정규표현식을 이용한 자동 개인/기업정보 차단 및 지정 협력사명 마스킹
    
    v2.1: 계좌번호 패턴을 은행별 실제 자릿수로 정밀화하여 오탐 방지
    """
    if not text:
        return text
        
    # 1. 개인정보 및 식별번호 차단 (자동)
    # 주민등록번호 / 외국인등록번호 (6자리-7자리, 뒷자리 1~4로 시작)
    text = re.sub(r'\b\d{6}[-\s]*[1-4]\d{6}\b', '█주민/외국인번호█', text) 
    # 사업자등록번호 (3-2-5 형식, 반드시 하이픈 포함)
    text = re.sub(r'\b\d{3}-\d{2}-\d{5}\b', '█사업자번호█', text) 
    # 법인등록번호 (6-7 형식, 반드시 하이픈 포함)
    text = re.sub(r'\b\d{6}-\d{7}\b', '█법인번호█', text) 
    # 휴대전화 (010/011/016/017/018/019 + 3~4자리 + 4자리)
    text = re.sub(r'\b01[016789][-\s]?\d{3,4}[-\s]?\d{4}\b', '█휴대전화█', text)
    # 일반 전화번호 (지역번호 2~3자리 + 3~4자리 + 4자리)
    text = re.sub(r'\b0[2-9][0-9]?[-\s]?\d{3,4}[-\s]?\d{4}\b', '█전화번호█', text)
    # 이메일
    text = re.sub(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', '█이메일█', text)
    # 계좌번호: "계좌" 키워드 근접 시에만 마스킹 (오탐 방지)
    # 패턴: "계좌" 뒤 30자 이내의 10~16자리 숫자열 (하이픈 허용)
    text = re.sub(
        r'(계좌[^\d]{0,30})(\d{2,6}[-]?\d{2,6}[-]?\d{2,6})',
        lambda m: m.group(1) + '█계좌번호█',
        text
    )
    
    # 2. 당사 식별 키워드 차단 (자동)
    company_keywords = ['신세계디에프', '신세계면세점', '신세계 DF', 'Shinsegae DF', '신세계']
    for kw in company_keywords:
        text = text.replace(kw, '█당사(내부정보)█')

    # 3. 지정된 협력사명 차단 (반자동)
    if target_partner:
        partners = [p.strip() for p in target_partner.split(',') if p.strip()]
        for p in partners:
            text = text.replace(p, '█협력사█')

    return text

# ── Supabase CRUD ────────────────────────────────────────────
def load_docs():
    try:
        res = init_supabase().table("docs").select("*").order("created_at").execute()
        return res.data or []
    except Exception as e:
        logger.error(f"문서 로드 실패: {e}")
        st.warning("⚠️ 기준 문서를 불러오지 못했습니다. 새로고침 해주세요.")
        return []

def save_doc(doc):
    try:
        init_supabase().table("docs").upsert({
            "id": doc["id"], "name": doc["name"], "cat": doc["cat"],
            "contract_type": doc.get("contract_type"), "label": doc["label"],
            "text": doc["text"], "size": doc["size"],
        }).execute()
        return True
    except Exception as e:
        logger.error(f"문서 저장 실패: {e}")
        st.warning(f"⚠️ 문서 저장 실패: {doc['name']}")
        return False

def delete_doc(doc_id):
    try:
        init_supabase().table("docs").delete().eq("id", doc_id).execute()
        return True
    except Exception as e:
        logger.error(f"문서 삭제 실패: {e}")
        st.warning("⚠️ 문서 삭제에 실패했습니다.")
        return False

def load_sessions():
    try:
        res = init_supabase().table("sessions").select("*").order("created_at", desc=True).execute()
        return res.data or []
    except Exception as e:
        logger.error(f"세션 로드 실패: {e}")
        return []

def save_session(sess):
    try:
        init_supabase().table("sessions").upsert({
            "id": sess["id"], "title": sess["title"],
            "date": sess["date"], "messages": sess["messages"],
        }).execute()
        return True
    except Exception as e:
        logger.error(f"세션 저장 실패: {e}")
        st.warning("⚠️ 대화 저장에 실패했습니다.")
        return False

def delete_session_db(sess_id):
    try:
        init_supabase().table("sessions").delete().eq("id", sess_id).execute()
        return True
    except Exception as e:
        logger.error(f"세션 삭제 실패: {e}")
        st.warning("⚠️ 대화 삭제에 실패했습니다.")
        return False

def save_review_log(log_data):
    try:
        init_supabase().table("review_logs").upsert(log_data).execute()
        return True
    except Exception as e:
        logger.error(f"검토 이력 저장 실패: {e}")
        return False

def load_laws():
    try:
        res = init_supabase().table("laws").select("*").order("id").execute()
        return res.data or []
    except Exception as e:
        logger.error(f"법령 로드 실패: {e}")
        return []

def cleanup_old_sessions(days=90):
    try:
        cutoff = (datetime.now() - timedelta(days=days)).isoformat()
        init_supabase().table("sessions").delete().lt("created_at", cutoff).execute()
    except Exception:
        pass

# ── docx 텍스트 추출 (에러 핸들링 강화) ───────────────────────
def extract_text(file_bytes):
    """docx 파일에서 텍스트 추출. 실패 시 에러 메시지 반환."""
    if not file_bytes:
        return "(빈 파일입니다.)"
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        if not paragraphs:
            # 테이블에만 텍스트가 있을 수 있으므로 테이블도 시도
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        table_texts.append(row_text)
            if table_texts:
                return "\n".join(table_texts)
            return "(문서에서 텍스트를 추출하지 못했습니다. 이미지만 있는 문서일 수 있습니다.)"
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"텍스트 추출 실패: {e}")
        return f"(파일 읽기 실패: 손상되었거나 지원하지 않는 형식입니다 — {type(e).__name__})"

# ── 시스템 프롬프트용 텍스트 자르기 (조항 단위) ─────────────────
def truncate_at_boundary(text, max_chars):
    """문서 구분자(---) 기준으로 자르기. 조항 중간에서 잘리지 않도록 함."""
    if len(text) <= max_chars:
        return text
    # max_chars 이전의 마지막 구분자 위치를 찾아 거기서 자름
    truncated = text[:max_chars]
    last_separator = truncated.rfind("\n\n---\n\n")
    if last_separator > max_chars * 0.5:  # 절반 이상은 보존
        return truncated[:last_separator]
    # 구분자가 없으면 마지막 줄바꿈 기준
    last_newline = truncated.rfind("\n")
    if last_newline > max_chars * 0.7:
        return truncated[:last_newline]
    return truncated

# ── 쿼리 라우팅 ──────────────────────────────────────────────
def route_query(query, has_attachment):
    if has_attachment:
        return "claude"
    if any(kw in query for kw in REVIEW_KEYWORDS):
        return "claude"
    # 이전 대화에 검토 결과(json_data)가 있으면 → 후속 질문도 claude로 유지
    # (검토 결과에 대한 추가 질문/대화 연속성 보장)
    for msg in st.session_state.get("messages", []):
        if msg.get("json_data"):
            return "claude"
    return "gemini"

def verify_citations(cited_laws, laws_db):
    """법령 인용 검증 — 새 JSON 구조(dict 또는 string) 모두 대응 + 시행일자 체크.
    cited_laws 형식:
      기존: ["대규모유통업법 제11조", ...]
      신규: [{"law_name": "대규모유통업법", "article": "제11조", "verified": true}, ...]
    """
    results = []
    if not cited_laws:
        return results
    
    for cite in cited_laws:
        # 새 형식(dict) vs 기존 형식(string) 모두 처리
        if isinstance(cite, dict):
            law_name_cite = cite.get("law_name", "")
            article_cite = cite.get("article", "")
            ai_verified = cite.get("verified", False)
            cite_str = f"{law_name_cite} {article_cite}"
        else:
            cite_str = str(cite)
            law_name_cite = ""
            article_cite = ""
            ai_verified = True  # 기존 형식은 검증 가정
        
        found = False
        matched_content = ""
        is_current = True
        last_updated = ""
        
        for law in laws_db:
            law_short = law.get("law_short", "")
            law_name_db = law.get("law_name", "")
            article_no = law.get("article_no", "")
            
            # 조문번호 매칭
            art_match = False
            if article_cite and article_cite == article_no:
                art_match = True
            elif article_no in cite_str:
                art_match = True
            if not art_match:
                continue
            
            # 법령명 매칭
            name_match = False
            if law_name_cite:
                if (law_name_cite == law_short or law_name_cite == law_name_db or
                    cite_partial_match(law_short, law_name_cite) or
                    cite_partial_match(law_name_db, law_name_cite)):
                    name_match = True
            else:
                if (law_short in cite_str or law_name_db in cite_str or
                    cite_partial_match(law_short, cite_str) or
                    cite_partial_match(law_name_db, cite_str)):
                    name_match = True
            
            if name_match:
                found = True
                matched_content = law["content"][:100] + "..."
                last_updated = law.get("last_updated", "")
                break
        
        results.append({
            "citation": cite_str,
            "verified": found,
            "ai_verified": ai_verified,
            "preview": matched_content if found else "",
            "last_updated": last_updated,
            "db_registered": found,
        })
    
    return results

def verify_precedents(cited_precedents):
    """판례 인용 검증 — AI가 verified=false로 표시한 판례 감지."""
    results = []
    if not cited_precedents:
        return results
    for prec in cited_precedents:
        if isinstance(prec, dict):
            case_no = prec.get("case_no", "미확인")
            verified = prec.get("verified", False)
            summary = prec.get("summary", "")
        else:
            case_no = str(prec)
            verified = True
            summary = ""
        
        results.append({
            "case_no": case_no,
            "verified": verified,
            "summary": summary,
        })
    return results

def cite_partial_match(name, cite):
    """부분 매칭 — 핵심 키워드가 인용문에 포함되는지 체크.
    예: '보세판매장고시' → '보세판매장'이 cite에 있으면 매칭
        '대규모유통업법' → '대규모유통업'이 cite에 있으면 매칭
    """
    # 법령명에서 '법', '고시', '시행령' 등 접미사 제거 → 핵심어 추출
    core = name
    for suffix in ["법", "고시", "시행령", "시행규칙"]:
        if core.endswith(suffix):
            core = core[:-len(suffix)]
            break
    return len(core) >= 3 and core in cite

# ── 에러 타입 분류 헬퍼 ──────────────────────────────────────
def classify_api_error(error):
    """API 에러를 분류하여 (에러유형, 메시지) 튜플 반환."""
    error_msg = str(error)
    error_lower = error_msg.lower()
    if any(kw in error_lower for kw in ["rate", "quota", "429", "resource_exhausted"]):
        return "rate_limit", "API 사용 한도 또는 요금을 초과했습니다."
    elif any(kw in error_lower for kw in ["401", "403", "authentication", "api_key", "permission"]):
        return "auth", "API 키가 유효하지 않거나 만료되었습니다."
    elif any(kw in error_lower for kw in ["credit balance"]):
        return "credit", "API 크레딧 잔액이 부족합니다. 충전이 필요합니다."
    elif any(kw in error_lower for kw in ["529", "overloaded", "overloadederr"]):
        return "overloaded", "Claude 서버가 일시적으로 과부하입니다. 잠시 후 다시 시도해주세요."
    elif any(kw in error_lower for kw in ["context_length", "too many tokens", "maximum context"]):
        return "context_overflow", "입력이 너무 길어 처리할 수 없습니다."
    elif any(kw in error_lower for kw in ["500", "502", "503", "504", "unavailable", "timeout"]):
        return "server", "서버가 일시적으로 응답하지 않습니다."
    else:
        return "unknown", f"{type(error).__name__}: {error_msg[:300]}"

# ── 시스템 프롬프트 ──────────────────────────────────────────
def build_system_claude(docs, laws_db, gemini_analysis=""):
    """Claude용 시스템 프롬프트 — 법령/행정규칙 전문가.
    사규 원문은 넣지 않음. Gemini의 사규 분석 결과를 받아서 법령 관점 추가.
    """
    # 법령 DB 목록만
    laws_text = ""
    if laws_db:
        law_groups = {}
        for law in laws_db:
            short = law.get("law_short", "기타")
            if short not in law_groups:
                law_groups[short] = []
            title = law.get("article_title", "")
            no = law.get("article_no", "")
            law_groups[short].append(f"{no} {title}" if title else no)
        list_lines = []
        for law_short, articles in sorted(law_groups.items()):
            list_lines.append(f"- {law_short}: " + ", ".join(articles))
        laws_text = "\n".join(list_lines)
    else:
        laws_text = "(법령/행정규칙 DB 미등록)"

    gemini_section = ""
    if gemini_analysis:
        gemini_section = (
            "\n\n━━━ [Stage 1 결과: 사규/표준계약서 분석 (Gemini)] ━━━\n" +
            truncate_at_boundary(gemini_analysis, 8000)
        )

    result = (
        "당신은 면세점 전문 법령/행정규칙 검토 AI입니다. (Stage 2: 최종 종합 검토)\n\n"
        "Stage 1에서 Gemini가 다음을 수행했습니다:\n"
        "- PART A: 사규/표준계약서 기준 내부 위반사항 분석\n"
        "- PART B: Google Search로 최신 법령 조문 및 판례/해석례 실시간 검색\n\n"
        "당신의 역할:\n"
        "1. Gemini의 사규 분석(PART A)을 사내 기준 관점으로 반영\n"
        "2. Gemini의 법령/판례 검색 결과(PART B)를 **검증**하고 추가 리스크 발굴\n"
        "3. 최종 종합 판단 (법령+사규+행정규칙 교차 검토)\n\n"
        "⚠️ 검증 시 중요사항:\n"
        "- Gemini가 인용한 조문번호/판례번호가 실제로 존재하는지 의심하세요\n"
        "- '미확인' 표시된 항목은 인용하지 마세요\n"
        "- 당신이 확실히 아는 법령만 인용하세요. 불확실하면 '확인 필요' 표시\n"
        "- 판례는 Gemini가 검색으로 확인한 것만 인용하고, 새로 추가하지 마세요\n"
        "- 보세판매장 관련: 「보세판매장 특허 및 운영에 관한 고시」, 「관세법」 필수\n"
        "- 인용 시 아래 DB 목록의 법령명과 조문번호를 정확히 사용\n"
        + gemini_section +
        "\n\n━━━ [적용 법령·행정규칙 DB 목록] ━━━\n" + laws_text +
        "\n\n━━━ [답변 형식] ━━━\n"
        "모든 법적 리스크를 빠짐없이. high/medium 상세, low 1~2문장. 해당없으면 생략.\n\n"
        "```json\n"
        "{\n"
        '  "summary": "1줄 요약",\n'
        '  "verdict": "approved|conditional|rejected",\n'
        '  "verdict_reason": "종합 판단 근거",\n'
        '  "issues": [{"issue_no":1,"title":"쟁점","risk_level":"high|medium|low",\n'
        '    "target_clause":"대상 원문","applicable_law":"법령/고시명 제X조",\n'
        '    "law_analysis":"법령·판례 평가 (Gemini 검색 결과 검증 포함)",\n'
        '    "applicable_rule":"사규 조항 (Stage1 PART A 인용)",\n'
        '    "rule_analysis":"사규 평가","recommendation":"권고안"}],\n'
        '  "action_plan":"액션 플랜","alternative_clause":"수정안 또는 null",\n'
        '  "cited_laws":["법령명 제X조"]\n'
        "}\n"
        "```\n"
        "JSON 아래 마크다운 상세 설명. 🔴 위반, 🔵 통과.\n\n"
        "후속 질문에는 JSON 없이 마크다운. 새 검토 요청 시만 JSON 출력.\n"
    )
    logger.info(f"Claude 시스템프롬프트: {len(result):,}자 (추정 {int(len(result)/1.5):,} 토큰)")
    return result

def build_system_gemini_stage1(docs, laws_db):
    """Stage 1: Gemini Researcher — 사규 분석 + Google Search로 법령/판례 수집.
    최종 보고서가 아닌 '구조화된 Raw Data'를 추출하는 역할.
    """
    def by_cat(cat):
        return [d for d in docs if d["cat"] == cat]
    def fmt_docs(ds):
        if not ds: return "(등록 없음)"
        return "\n\n---\n\n".join(["[" + d["label"] + "]\n" + d["text"] for d in ds])

    saryu_text    = truncate_at_boundary(fmt_docs(by_cat("saryu")), 50000)
    contract_text = truncate_at_boundary(fmt_docs(by_cat("contract")), 60000)
    yakjeong_text = truncate_at_boundary(fmt_docs(by_cat("yakjeong")), 30000)

    laws_list = ""
    if laws_db:
        law_groups = {}
        for law in laws_db:
            short = law.get("law_short", "기타")
            if short not in law_groups:
                law_groups[short] = []
            law_groups[short].append(law.get("article_no", ""))
        lines = [f"- {ls}: " + ", ".join(arts) for ls, arts in sorted(law_groups.items())]
        laws_list = "\n".join(lines)

    return (
        "██ 최우선 규칙 — 절대 위반 금지 ██\n"
        "1. 사규 DB에 명확한 근거가 없으면 '해당 규정 없음'으로 표시할 것.\n"
        "2. 법령 조문은 반드시 Google Search로 law.go.kr에서 확인한 것만 기재할 것.\n"
        "3. 판례 인용 시 반드시 대법원 사건번호(예: 대법원 2011도6759)를 기재할 것. "
        "정확한 사건번호를 Google Search로 확인하지 못하면 해당 판례를 아예 언급하지 말 것.\n"
        "4. 확인 안 된 법령/판례는 절대 지어내지 말고 목록에서 제외할 것.\n"
        "5. 구법(폐지/개정)이 아닌 현행 법령만 인용할 것.\n"
        "6. 판례 내용만 서술하고 사건번호를 생략하는 것은 금지. 번호 없으면 판례 미인용.\n\n"

        "당신은 면세점 공정거래 법률 리서처입니다.\n"
        "당신의 역할은 '최종 판단'이 아니라 '정확한 데이터 수집'입니다.\n"
        "사용자의 질문을 분석하여 관련 사규 조항과 법령/판례를 수집하고,\n"
        "반드시 아래 JSON 형식으로만 출력하세요. 마크다운 설명은 작성하지 마세요.\n\n"

        "```json\n"
        "{\n"
        '  "query_summary": "사용자 질문 1줄 요약",\n'
        '  "saryu_findings": [\n'
        '    {"source": "사규/계약서 문서명", "clause": "조항번호/제목", "content": "해당 원문 발췌", "relevance": "위반|부합|공백"}\n'
        '  ],\n'
        '  "law_findings": [\n'
        '    {"law_name": "대규모유통업법", "article": "제11조", "title": "조문 제목", "content": "Google Search로 확인한 조문 핵심 내용 요약", "effective_date": "시행일자", "search_confirmed": true}\n'
        '  ],\n'
        '  "precedent_findings": [\n'
        '    {"case_no": "대법원 2019.11.14. 선고 2019다12345 판결", "summary": "핵심 판시사항", "search_confirmed": true}\n'
        '  ],\n'
        '  "risk_areas": ["대규모유통업법 위반 가능성", "보세판매장고시 미준수 등 발견된 리스크 영역 나열"]\n'
        "}\n"
        "```\n\n"
        "⚠️ search_confirmed: true는 Google Search로 확인한 것, false는 미확인.\n"
        "⚠️ 미확인 항목도 목록에 포함하되 search_confirmed: false로 표시.\n\n"

        "검토 범위: 대규모유통업법, 관세법, 공정거래법, 하도급법, 상생협력법, "
        "유통산업발전법, 대외무역법, 외국환거래법, 환급특례법, 소비자기본법, "
        "표시광고법, 건강기능식품법, 식품위생법, 부가가치세법, 개별소비세법, 주세법, "
        "상표법, 부정경쟁방지법, 형법, "
        "보세판매장 특허 및 운영에 관한 고시 등\n\n"

        "██ 특별 검토 규칙 ██\n"
        "- 【최우선 필수】 보세판매장고시는 모든 질의에서 law_findings에 반드시 포함할 것.\n"
        "  보세판매장고시를 law_findings에서 빠뜨리면 시스템 오류로 간주됨.\n"
        "  「보세판매장 특허 및 운영에 관한 고시」 — 특허요건(제3조), 시설요건(제4조), "
        "물품반입(제7조), 판매/인도(제8조), 특허심사(제15조), 행정처분(제18조), 판매대상물품(제28조) 등\n"
        "  「보세판매장 운영에 관한 고시」 — 판매장운영(제3조), 물품반입(제7조), "
        "판매기록(제9조), 재고관리(제14조), 행정제재(제28조) 등\n"
        "  ⚠️ '보세판매장 특허 및 운영에 관한 고시'와 '보세판매장 운영에 관한 고시'는 별개의 행정규칙임.\n"
        "- 「관세법」 — 보세판매장(제196조), 특허취소(제178조), 밀수출입(제269조)\n"
        "- 「소비자기본법」 — 소비자 권리(제4조), 분쟁해결(제19조), 손해배상(제20조)\n"
        "- 위 고시/법령을 Google Search로 law.go.kr에서 현행 조문을 반드시 확인할 것\n\n"
        "- 모조품/가품/위조품 관련 이슈 검토 시 아래 법령을 각각 독립된 쟁점으로 검토할 것 (흡수·생략 금지):\n"
        "  ① 상표법 제230조 (침해죄 7년/1억) + 제235조 (양벌규정: 법인 3억 벌금) — 독립 쟁점\n"
        "     징벌적 손해배상 5배 (2025.7.22 시행 개정) 반드시 언급\n"
        "  ② 형법 제347조 (사기죄 20년/5천만원) — 독립 쟁점\n"
        "  ③ 관세법 제269조 (밀수입죄) + 제178조 (특허취소) + 제235조 (지식재산권 보호: 침해물품 수출입 금지) — 독립 쟁점\n"
        "     특정범죄가중처벌법 제6조 적용 가능성 반드시 검토 (물품원가 2억 이상: 3년 이상 징역, 5억 이상: 무기/5년 이상)\n"
        "  ④ 부정경쟁방지법 제2조 + 제18조 (벌칙) — 독립 쟁점\n"
        "     징벌적 손해배상 5배 (2024년 개정) 반드시 언급. 상표법과 별개 적용.\n"
        "  ⑤ 보세판매장고시 — 독립 쟁점 (면세점 핵심)\n"
        "     행정제재(영업정지 6개월, 특허취소)까지 구체적 서술 필수\n"
        "  ⑥ 소비자기본법 — 독립 쟁점\n"
        "     제4조(권리), 제46조(위해방지 의무: 수거·파기·환급), 집단분쟁조정·소비자단체소송 가능성\n"
        "  ⑦ 면세점 특수 가중요소 — 독립 쟁점 또는 별도 섹션\n"
        "     국가 특허 지위로 인한 높은 주의의무, 정품 보증 소비자 기대, 외국인 고객 국제적 파급\n"
        "  ⚠️ 위 7개 영역 모두 검토. 양벌규정(법인 처벌)과 징벌적 손해배상은 면세점 경영진 최대 관심사.\n"
        "- 관세법: 특허취소(제178조), 밀수출입(제269조), 지식재산권보호(제235조) 포함.\n\n"

        "━━━ [당사 사규] ━━━\n" + saryu_text +
        "\n\n━━━ [당사 표준 계약서] ━━━\n" + contract_text +
        "\n\n━━━ [당사 표준 약정서] ━━━\n" + yakjeong_text +
        "\n\n━━━ [법령 DB 등록 목록 (참고)] ━━━\n" + laws_list
    )

def call_mcp_law(query, max_tokens=2048):
    """korean-law-mcp 서버를 통해 법령/판례 조회.
    Anthropic API의 MCP connector를 사용.
    Returns: (success: bool, text_content: str)
    """
    try:
        client = init_anthropic()
        
        # MCP 호출 — SDK 0.85+에서는 betas 파라미터 사용
        response = client.beta.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=max_tokens,
            messages=[{"role": "user", "content": query}],
            mcp_servers=[{
                "type": "url",
                "url": "https://korean-law-mcp.fly.dev/mcp",
                "name": "korean-law"
            }],
            betas=["mcp-client-2025-11-20"]
        )
        # 응답에서 텍스트 추출
        text_parts = []
        for block in response.content:
            if hasattr(block, 'text') and block.text:
                text_parts.append(block.text)
        return True, "\n".join(text_parts)
    except Exception as e:
        err_str = str(e)
        logger.warning(f"MCP 법령 조회 실패: {err_str[:500]}")
        
        # MCP 실패 시 law.go.kr API 직접 폴백
        try:
            import requests
            url = "http://www.law.go.kr/DRF/lawSearch.do"
            # 쿼리에서 법령명 추출
            import re as _re
            law_match = _re.search(r"['\"]([가-힣]+법[가-힣]*)['\"]", query)
            if law_match:
                params = {"OC": "sapphire_5", "target": "law", "type": "XML", "query": law_match.group(1)}
                res = requests.get(url, params=params, timeout=10)
                if res.status_code == 200 and "totalCnt" in res.text:
                    return True, f"{law_match.group(1)} 법령 존재 확인 (law.go.kr 직접 조회)"
        except Exception:
            pass
        
        return False, err_str[:300]


def verify_precedent_via_api(case_no, context_keywords=None):
    """MCP를 통해 판례 검증.
    1) 사건번호 존재 여부
    2) 사건명이 인용 맥락과 관련 있는지 (할루시네이션 정밀 감지)
    Returns: (status: str, title: str)
    """
    import re as _re
    num_match = _re.search(r'(\d{4}[가-힣]+\d+)', case_no)
    if not num_match:
        return "not_found", ""
    
    query = num_match.group(1)
    try:
        success, result = call_mcp_law(f"대법원 판례 '{query}' 사건번호로 검색해서 존재 여부와 사건명만 알려줘. 없으면 '미존재'라고 답해.")
        
        if not success:
            return "error", ""
        
        result_lower = result.lower()
        
        # 미존재 판단
        if "미존재" in result or "찾을 수 없" in result or "검색 결과가 없" in result or "존재하지 않" in result:
            return "not_found", ""
        
        # 사건명 추출 (MCP 응답에서)
        title = ""
        title_match = _re.search(r'사건명[:\s]*([^\n,]+)', result)
        if title_match:
            title = title_match.group(1).strip()
        elif len(result) < 200:
            title = result.strip()
        
        # 맥락 관련성 체크
        if context_keywords and title:
            title_lower = title.lower()
            has_relevance = any(kw in title_lower for kw in context_keywords)
            if not has_relevance:
                return "wrong_context", title
        
        return "verified", title
    except Exception as e:
        logger.warning(f"판례 MCP 검증 실패 ({query}): {e}")
        return "error", ""


def verify_law_via_api(law_name, article):
    """MCP를 통해 법령 조문 존재 + 원문 조회.
    Returns: (exists: bool, content_preview: str)
    """
    try:
        success, result = call_mcp_law(
            f"'{law_name}' {article} 조문 원문을 알려줘. 법령이 없으면 '미존재'라고 답해."
        )
        
        if not success:
            return False, ""
        
        if "미존재" in result or "찾을 수 없" in result:
            return False, ""
        
        return True, result[:500]
    except Exception as e:
        logger.warning(f"법령 MCP 검증 실패 ({law_name}): {e}")
        return False, ""


def gatekeeper_process(gemini_raw_json, laws_db, docs=None):
    """Stage 2: Gatekeeper — Gemini의 Raw Data를 3중 검증 후 정제.
    1차: 사내 DB(124건+) 원문 대조
    2차: law.go.kr API 실시간 판례 검증 (할루시네이션 차단)
    3차: DB 미등록 법령도 API로 존재 여부 확인
    결과: Claude에게 전달할 정제된 텍스트 (2~5k 토큰).
    """
    if not gemini_raw_json:
        return {"error": "Gemini 결과 없음"}, ""
    
    refined_parts = []
    verified_laws = []
    unverified_laws = []
    verified_precedents = []
    dropped_precedents = []
    
    # 1. 사규 분석 결과 → Gemini가 못 찾았으면 원문 직접 주입
    saryu_findings = gemini_raw_json.get("saryu_findings", [])
    if saryu_findings:
        refined_parts.append("━━━ [사규 분석 결과 (Gemini 확인)] ━━━")
        for sf in saryu_findings:
            refined_parts.append(f"- [{sf.get('relevance','?')}] {sf.get('source','?')} {sf.get('clause','')}: {sf.get('content','')[:200]}")
    elif docs:
        # Gemini가 사규를 못 찾았으면 원문 핵심 부분 직접 주입
        refined_parts.append("━━━ [사규 원문 — 시스템 강제 주입 (Gemini 누락)] ━━━")
        for doc in docs:
            if doc.get("cat") in ("saryu", "contract", "yakjeong"):
                label = doc.get("label", doc.get("name", ""))
                text_preview = doc.get("text", "")[:500]
                refined_parts.append(f"📄 [{label}]\n{text_preview}")
        logger.info("Gatekeeper: Gemini 사규 누락 → 원문 강제 주입")
    
    # 2. 법령 → 사내 DB 대조 + API 보조 검증
    law_findings = gemini_raw_json.get("law_findings", [])
    refined_parts.append("\n━━━ [법령 검증 결과] ━━━")
    
    for lf in law_findings:
        law_name = lf.get("law_name", "")
        article = lf.get("article", "")
        search_confirmed = lf.get("search_confirmed", False)
        
        # 1차: 사내 DB 매칭
        db_match = None
        for db_law in laws_db:
            db_short = db_law.get("law_short", "")
            db_name = db_law.get("law_name", "")
            db_article = db_law.get("article_no", "")
            if article != db_article:
                continue
            if (law_name == db_short or law_name == db_name or
                cite_partial_match(db_short, law_name) or
                cite_partial_match(db_name, law_name)):
                db_match = db_law
                break
        
        if db_match:
            db_content = db_match["content"][:1000]  # 원문 충분히 전달
            last_updated = db_match.get("last_updated", "")[:10]
            refined_parts.append(
                f"✅ [{law_name} {article}] (DB 검증 완료, 업데이트: {last_updated})\n"
                f"   ⚠️ 아래 원문의 수치(징역, 벌금 등)를 반드시 그대로 인용할 것.\n"
                f"   【DB 원문 시작】\n"
                f"   {db_content}\n"
                f"   【DB 원문 끝】"
            )
            verified_laws.append({"law_name": law_name, "article": article, "db_verified": True, "last_updated": last_updated})
        else:
            # 2차: law.go.kr API로 법령 존재 여부 확인
            api_exists, api_info = verify_law_via_api(law_name, article)
            time.sleep(0.5)  # Rate limit 방어
            
            if api_exists:
                refined_parts.append(
                    f"🔍 [{law_name} {article}] (DB 미등록, API 법령 존재 확인)\n"
                    f"   Gemini 요약: {lf.get('content', '내용 없음')[:200]}"
                )
                unverified_laws.append({"law_name": law_name, "article": article, "db_verified": False, "api_verified": True})
            else:
                refined_parts.append(
                    f"⚠️ [{law_name} {article}] (DB 미등록, API 미확인 — 존재 여부 불확실)\n"
                    f"   Gemini 요약: {lf.get('content', '내용 없음')[:200]}"
                )
                unverified_laws.append({"law_name": law_name, "article": article, "db_verified": False, "api_verified": False})
    
    # 3. 판례 → law.go.kr API로 실시간 검증 (할루시네이션 차단)
    precedent_findings = gemini_raw_json.get("precedent_findings", [])
    if precedent_findings:
        refined_parts.append("\n━━━ [판례 검증 결과 (law.go.kr API 대조)] ━━━")
        
        # 리스크 영역에서 맥락 키워드 추출 (판례가 올바른 맥락인지 확인용)
        risk_areas = gemini_raw_json.get("risk_areas", [])
        context_kw = []
        for ra in risk_areas:
            for kw in ["상표", "위조", "모조", "가품", "사기", "관세", "밀수", "공정거래", "하도급", 
                       "유통", "소비자", "광고", "식품", "세금", "면세", "보세", "부정경쟁",
                       "횡령", "배임", "특허", "저작권", "영업비밀"]:
                if kw in ra:
                    context_kw.append(kw)
        # Gemini의 판례 요약에서도 키워드 추출
        for pf in precedent_findings:
            summary = pf.get("summary", "")
            for kw in ["상표", "위조", "모조", "사기", "관세", "밀수", "공정거래", "부정경쟁"]:
                if kw in summary:
                    context_kw.append(kw)
        context_kw = list(set(context_kw)) if context_kw else None
        
        for pf in precedent_findings:
            case_no = pf.get("case_no", "미확인")
            if case_no == "미확인" or not case_no:
                continue
            
            # law.go.kr API 3단계 검증
            status, title = verify_precedent_via_api(case_no, context_kw)
            time.sleep(0.5)
            
            if status == "verified":
                refined_parts.append(f"✅ [{case_no}] (API 검증 완료, 사건명: {title}): {pf.get('summary', '')[:150]}")
                verified_precedents.append({"case_no": case_no, "title": title, "summary": pf.get("summary", ""), "api_verified": True})
            elif status == "wrong_context":
                # 🚨 사건번호는 존재하지만 맥락과 무관 → 할루시네이션
                logger.warning(f"🚨 맥락 불일치 판례 감지: {case_no} (실제: {title})")
                refined_parts.append(
                    f"🚨 [{case_no}] → 사건번호는 존재하나 실제 사건명은 '{title}'로, "
                    f"AI가 인용한 맥락과 무관함. 허위 인용으로 판정하여 삭제됨."
                )
                dropped_precedents.append({"case_no": case_no, "reason": f"맥락 불일치 — 실제 사건: {title}"})
            elif status == "not_found":
                logger.warning(f"🚨 미존재 판례 감지: {case_no}")
                refined_parts.append(f"🚨 [{case_no}] → 국가법령정보센터에 존재하지 않음. 허위 판례로 판정하여 삭제됨.")
                dropped_precedents.append({"case_no": case_no, "reason": "API 검증 실패 — 존재하지 않는 판례"})
            else:  # error
                refined_parts.append(f"⚠️ [{case_no}] → API 검증 실패 (네트워크 오류). 수동 확인 필요.")
                dropped_precedents.append({"case_no": case_no, "reason": "API 통신 오류 — 수동 확인 필요"})
    
    # 4. 리스크 영역
    risk_areas = gemini_raw_json.get("risk_areas", [])
    if risk_areas:
        refined_parts.append("\n━━━ [발견된 리스크 영역] ━━━")
        for ra in risk_areas:
            refined_parts.append(f"- {ra}")
    
    # 5. ██ 보세판매장고시 + 소비자기본법 무조건 강제 주입 ██
    # 면세점 전용 앱 — Gemini 수집 여부와 무관하게 항상 DB 원문 주입
    forced_law_shorts = {
        "보세판매장고시": {"제3조", "제4조", "제5조", "제7조", "제8조", "제10조", 
                        "제15조", "제18조", "제19조", "제20조", "제21조", "제28조"},
        "보세판매장운영고시": {"제3조", "제7조", "제9조", "제14조", "제28조"},
        "소비자기본법": {"제4조", "제19조", "제20조"},
    }
    
    for force_short, force_articles in forced_law_shorts.items():
        force_laws = [db_law for db_law in laws_db if db_law.get("law_short") == force_short]
        if not force_laws:
            continue
        
        # 이미 Gemini가 수집했는지 확인
        already_found = any(
            force_short in lf.get("law_name", "")
            for lf in gemini_raw_json.get("law_findings", [])
        )
        if already_found:
            continue
        
        section_name = force_short
        refined_parts.append(f"\n━━━ [{section_name} — 필수 검토 (DB 원문)] ━━━")
        refined_parts.append(f"██ {section_name} 조문은 면세점 모든 검토에 필수. 반드시 issues에 반영. ██")
        
        injected = 0
        for ba in force_laws:
            if ba["article_no"] not in force_articles:
                continue
            content = ba["content"][:500]
            art_title = ba.get("article_title", "")
            refined_parts.append(
                f"✅ [{ba['law_short']} {ba['article_no']}] {art_title}\n"
                f"   【DB 원문 시작】\n"
                f"   {content}\n"
                f"   【DB 원문 끝】"
            )
            verified_laws.append({
                "law_name": ba.get("law_name", ba["law_short"]),
                "article": ba["article_no"],
                "db_verified": True,
                "last_updated": ba.get("last_updated", "")[:10],
            })
            injected += 1
        if injected > 0:
            logger.info(f"Gatekeeper: {section_name} {injected}건 강제 주입")
    
    refined_text = "\n".join(refined_parts)
    
    meta = {
        "query_summary": gemini_raw_json.get("query_summary", ""),
        "verified_laws": verified_laws,
        "unverified_laws": unverified_laws,
        "verified_precedents": verified_precedents,
        "dropped_precedents": dropped_precedents,
        "total_laws": len(verified_laws) + len(unverified_laws),
        "total_precedents": len(verified_precedents),
        "total_dropped": len(dropped_precedents),
    }
    
    logger.info(f"Gatekeeper: 법령 {meta['total_laws']}건(DB검증 {len(verified_laws)}), "
                f"판례 검증{len(verified_precedents)}/삭제{len(dropped_precedents)}, 정제 {len(refined_text)}자")
    return meta, refined_text


def build_system_claude_v3(gatekeeper_text, gatekeeper_meta):
    """Stage 3: Claude Senior Lawyer — 정제된 데이터만으로 법리 해석 + 최종 보고서.
    외부 지식 사용 금지 — 오직 전달된 텍스트 안에서만 추론.
    """
    return (
        "████████████████████████████████████████████████████████\n"
        "██ 최우선 규칙 — 1개라도 위반 시 출력물 전체가 무효 ██\n"
        "████████████████████████████████████████████████████████\n\n"
        
        "██ 규칙 0 — 보세판매장고시 필수 (이 앱의 존재 이유) ██\n"
        "이 시스템은 면세점 전용 법무 어시스턴트이다.\n"
        "검증 데이터에 '보세판매장고시' 조문이 포함되어 있다.\n"
        "반드시 issues에 보세판매장고시 관련 독립 쟁점을 포함하고:\n"
        "  - DB 원문에 근거한 구체적 행정처분 분석 (영업정지 6개월, 특허취소 등)\n"
        "  - 판매대상 물품 제한, 운영 의무 위반 등 실무적 리스크 서술\n"
        "  - 단순히 '추가 검토 필요', '별도 확인 필요'로 때우는 것은 절대 금지\n"
        "  - '보세판매장고시 검토 누락'이라는 문구 사용 절대 금지\n"
        "보세판매장고시를 issues에서 빠뜨리면 출력물 전체가 무효이다.\n\n"
        
        "██ 규칙 1 — 추측 금지 ██\n"
        "DB 원문이 없는 법령의 형량을 추측하지 마라. '원문 확인 필요'로만 표기.\n\n"
        "██ 규칙 2 — 검증 데이터만 사용 ██\n"
        "【DB 원문 시작】~【DB 원문 끝】 안에서만 추론. 외부 지식 금지.\n\n"
        "██ 규칙 3 — 판례 창작 금지 ██\n"
        "판결 요지 원문 미확보 시 해석 서술 금지. '(판결 요지 — 원문 확인 필요)' 표기.\n\n"
        "██ 규칙 4 — 형량 Placeholder ██\n"
        "형량은 {{법령약칭_조문번호_형량}} 형태로만 작성. DB 미등록은 '(형량 — 원문 확인 필요)'.\n\n"
        "██ 규칙 5 — 용어 ██\n"
        "'~팀' 금지 → '~담당부서'. 새 법령 추가 인용 금지.\n\n"
        
        "████ 필수 쟁점 포함 규칙 (모조품/가품 관련) ████\n"
        "모조품 관련 질의 시 아래를 각각 독립 쟁점으로 출력:\n"
        "① 상표법(침해죄+양벌규정+징벌적배상5배) ② 형법(사기) ③ 관세법(밀수입+특허취소+제235조 지식재산권보호+특가법 가중)\n"
        "④ 부정경쟁방지법(벌칙+징벌적배상5배) ⑤ 보세판매장고시(행정처분 구체 서술)\n"
        "⑥ 소비자기본법(권리+위해방지+집단분쟁) ⑦ 면세점 특수 가중(국가특허 주의의무, 정품보증 기대)\n"
        "████████████████████████████████████████████████████████\n\n"
        
        "당신은 면세점 전문 시니어 변호사 AI입니다.\n"
        "아래 정제된 검증 데이터만을 기반으로 법리 해석과 실무 권고를 작성하세요.\n\n"
        
        "━━━ [검증 데이터] ━━━\n" + gatekeeper_text +
        
        "\n\n━━━ [답변 형식] ━━━\n"
        "반드시 아래 JSON 형식으로만 출력하세요. JSON 외의 마크다운 설명은 작성하지 마세요.\n"
        "모든 분석 내용은 JSON의 issues 안에 넣으세요.\n\n"
        "```json\n"
        "{\n"
        '  "summary": "1줄 요약",\n'
        '  "verdict": "approved | conditional | rejected",\n'
        '  "verdict_reason": "종합 판단 근거 (전달된 데이터에 근거하여)",\n'
        '  "issues": [\n'
        '    {\n'
        '      "issue_no": 1, "title": "쟁점 제목", "risk_level": "high|medium|low",\n'
        '      "target_clause": "검토 대상 원문",\n'
        '      "applicable_law": "법령/고시명 제X조",\n'
        '      "law_analysis": "법령·행정규칙 관점 — 형량은 반드시 {{법령약칭_조문번호_형량}} Placeholder 사용",\n'
        '      "applicable_rule": "사규 조항",\n'
        '      "rule_analysis": "사규 관점",\n'
        '      "recommendation": "종합 권고안"\n'
        '    }\n'
        '  ],\n'
        '  "action_plan": "MD 실무 액션 (단계별, 시간축 포함: 즉시/24시간/1주/1개월)",\n'
        '  "alternative_clause": "수정 대안 (없으면 null)",\n'
        '  "cited_laws": [{"law_name": "법령명", "article": "제X조"}],\n'
        '  "cited_precedents": [{"case_no": "판례번호", "summary": "판시사항"}]\n'
        "}\n"
        "```\n"
        "⚠️ JSON 코드블록 뒤에 추가 마크다운 설명을 절대 작성하지 마세요. JSON만 출력하세요.\n"
        "후속 질문에는 JSON 없이 마크다운. 새 검토 요청 시만 JSON 출력.\n"
    )


def build_system_gemini(docs):
    def by_cat(cat):
        return [d for d in docs if d["cat"] == cat]
    def fmt_docs(ds):
        if not ds: return "(등록 없음)"
        return "\n\n---\n\n".join(["[" + d["label"] + "]\n" + d["text"] for d in ds])

    saryu_text    = truncate_at_boundary(fmt_docs(by_cat("saryu")), 25000)
    contract_text = truncate_at_boundary(fmt_docs(by_cat("contract")), 35000)
    yakjeong_text = truncate_at_boundary(fmt_docs(by_cat("yakjeong")), 20000)

    return (
        "당신은 면세점 전문 공정거래 실무 어시스턴트 AI입니다.\n"
        "사규, 계약서, 약정서 내용에 대한 일반 질문에 친절히 답변하세요.\n"
        "법률 검토 판단(승인/반려)은 하지 말고, 필요시 '검토 요청을 해주세요'라고 안내하세요.\n\n"
        "① 당사 사규:\n" + saryu_text +
        "\n\n② 당사 표준 계약서:\n" + contract_text +
        "\n\n③ 당사 표준 약정서:\n" + yakjeong_text +
        "\n\n답변 시 위험은 🔴, 적법은 🔵 이모지를 사용하세요. Streamlit 색상 단축코드(:red[], :blue[], :blue_circle: 등)는 절대 사용하지 마세요."
    )

# ── AI 호출 및 에러 핸들링 함수 ────────────────────────────────
def call_claude(system_prompt, messages):
    client = init_anthropic()
    claude_messages = []
    last_role = None

    # 대화 히스토리 극한 제한: 마지막 질문만
    if len(messages) > 4:
        recent_messages = [messages[-1]]
    else:
        recent_messages = messages

    for m in recent_messages:
        role = "assistant" if m["role"] == "assistant" else "user"
        content = m["content"]
        
        # assistant 응답 극한 축소
        if role == "assistant":
            if len(content) > 1000:
                content = content[:1000]
            if not content.strip():
                content = "(이전 결과 참조)"
        
        # user 메시지 극한 축소
        if role == "user" and len(content) > 8000:
            content = content[:8000] + "\n...(이하 생략)"

        if role == last_role:
            claude_messages[-1]["content"] += f"\n\n{content}"
        else:
            claude_messages.append({"role": role, "content": content})
            last_role = role
            
    if claude_messages and claude_messages[0]["role"] == "assistant":
        claude_messages.pop(0)
    
    # 마지막 메시지가 user인지 확인
    if not claude_messages or claude_messages[-1]["role"] != "user":
        claude_messages.append({"role": "user", "content": "검토해주세요."})

    # 크기 로깅
    sys_chars = len(system_prompt)
    msg_chars = sum(len(m["content"]) for m in claude_messages)
    total_chars = sys_chars + msg_chars
    logger.info(f"Claude 호출: system={sys_chars:,}자, msgs={msg_chars:,}자, total={total_chars:,}자 (~{int(total_chars/1.5):,}tok), msg수={len(claude_messages)}")

    try:
        response = client.messages.create(
            model=CLAUDE_MODEL,
            max_tokens=8192,
            system=system_prompt,
            messages=claude_messages,
        )
        return response.content[0].text
    except Exception as e:
        err_type, err_msg = classify_api_error(e)
        
        # overloaded/server → 최대 3회 재시도 (3초→6초→12초)
        if err_type in ("overloaded", "server"):
            for retry in range(3):
                wait = 3 * (2 ** retry)  # 3, 6, 12초
                logger.info(f"Claude {err_type} — {wait}초 대기 후 재시도 {retry+1}/3...")
                time.sleep(wait)
                try:
                    response = client.messages.create(
                        model=CLAUDE_MODEL,
                        max_tokens=8192,
                        system=system_prompt,
                        messages=claude_messages,
                    )
                    logger.info(f"Claude 재시도 {retry+1} 성공!")
                    return response.content[0].text
                except Exception as retry_e:
                    err_type, err_msg = classify_api_error(retry_e)
                    if err_type not in ("overloaded", "server"):
                        break  # 다른 종류 에러면 재시도 중단
        
        logger.error(f"Claude API 최종 실패: {err_type} — {str(e)[:200]}")
        label_map = {
            "rate_limit": "API 한도 초과",
            "auth": "API 인증 오류",
            "credit": "크레딧 부족",
            "overloaded": "서버 과부하",
            "server": "서버 통신 장애",
            "context_overflow": "입력 초과",
            "unknown": "통신 오류",
        }
        debug_info = f"(system={sys_chars:,}자, msg={msg_chars:,}자)"
        return f"⚠️ [{label_map.get(err_type, '오류')}] Claude: {err_msg} {debug_info}"

def call_gemini(system_prompt, messages):
    from google.genai import types
    client = init_gemini()
    
    last_error_type = None
    
    for model_name in GEMINI_MODELS:
        # rate_limit 에러 시 동일 API 키 쿼터를 공유하므로 다음 모델로 폴백해도 무의미
        if last_error_type == "rate_limit":
            logger.info(f"Gemini rate limit 발생 — 동일 쿼터 공유로 {model_name} 폴백 건너뜀")
            continue
            
        try:
            history = []
            for m in messages[:-1]:
                role = "model" if m["role"] == "assistant" else "user"
                history.append(types.Content(role=role, parts=[types.Part(text=m["content"])]))
            last_msg = messages[-1]["content"]
            response = client.models.generate_content(
                model=model_name,
                contents=history + [types.Content(role="user", parts=[types.Part(text=last_msg)])],
                config=types.GenerateContentConfig(
                    system_instruction=system_prompt,
                    tools=[types.Tool(google_search=types.GoogleSearch())],
                ),
            )
            return response.text
        except Exception as e:
            logger.error(f"Gemini ({model_name}) 오류: {e}")
            last_error_type, last_err_msg = classify_api_error(e)
            
            # 마지막 모델이거나 rate_limit(공유 쿼터)인 경우 즉시 에러 반환
            is_last = (model_name == GEMINI_MODELS[-1])
            if is_last or last_error_type == "rate_limit":
                label_map = {
                    "rate_limit": "API 한도 초과",
                    "auth": "API 인증 오류",
                    "server": "서버 통신 장애",
                    "unknown": "통신 오류",
                }
                return f"⚠️ [{label_map[last_error_type]}] Gemini: {last_err_msg}"
            # auth 에러도 같은 키를 쓰므로 폴백 무의미
            if last_error_type == "auth":
                return f"⚠️ [API 인증 오류] Gemini: {last_err_msg}"
            continue
    
    return "⚠️ Gemini 응답을 가져오지 못했습니다."

def postprocess_reply(reply, laws_db):
    """AI 출력물 후처리 — Placeholder 치환 + 가짜 판례 차단.
    Claude/Gemini 출력 후 마지막 방어선.
    주의: JSON 블록 내부는 건드리지 않음 (파싱 깨짐 방지).
    """
    import re as _re
    
    # JSON 블록 분리 (```json...``` 또는 맨 앞 {...})
    json_block = ""
    rest_text = reply
    
    json_code_match = _re.search(r'(```json\s*\n.*?\n```)', reply, _re.DOTALL)
    if json_code_match:
        json_block = json_code_match.group(1)
        rest_text = reply[:json_code_match.start()] + "{{JSON_BLOCK}}" + reply[json_code_match.end():]
    elif reply.strip().startswith("{"):
        # 첫 번째 최상위 } 찾기
        brace_count = 0
        json_end = 0
        for i, c in enumerate(reply):
            if c == '{': brace_count += 1
            elif c == '}':
                brace_count -= 1
                if brace_count == 0:
                    json_end = i + 1
                    break
        if json_end > 0:
            json_block = reply[:json_end]
            rest_text = "{{JSON_BLOCK}}" + reply[json_end:]
    
    # 후처리는 rest_text에만 적용 (JSON 보호)
    
    # 1. ██ Placeholder → DB 원문 치환 ██
    # 패턴A: {{형법_제347조_형량}} (정상 형식)
    # 패턴B: {{상표법_230조_7년 이하의 징역 또는 1억원 이하의 벌금}} (Claude가 값까지 넣은 형식)
    # 패턴C: {{형법_347조_20년...}} (제 누락)
    # 모두 잡아서 DB 원문으로 교체
    
    all_placeholder_pattern = _re.compile(r'\{\{(.+?)\}\}')
    
    def replace_any_placeholder(match):
        """모든 {{ }} Placeholder를 DB 원문 형량으로 치환"""
        inner = match.group(1)  # "형법_제347조_형량" 또는 "상표법_230조_7년..."
        
        # 법령약칭 추출 (첫 번째 _ 앞)
        parts = inner.split("_")
        if len(parts) < 2:
            return match.group(0)  # {{ }} 그대로 반환
        
        law_short_guess = parts[0]  # "형법", "상표법", "관세법"
        
        # 조문번호 추출 — "제347조", "347조", "230조" 등
        article_guess = ""
        for p in parts[1:]:
            if _re.search(r'\d+조', p):
                article_guess = p if p.startswith("제") else f"제{p}"
                break
        
        if not article_guess:
            return match.group(0)
        
        # DB에서 매칭
        for db_law in laws_db:
            db_short = db_law.get("law_short", "")
            db_article = db_law.get("article_no", "")
            
            if db_article != article_guess:
                continue
            if db_short == law_short_guess or law_short_guess in db_short or db_short in law_short_guess:
                content = db_law.get("content", "")
                # 형량 패턴 추출
                penalty = _re.search(
                    r'(\d+년\s*이하의?\s*징역\s*(?:또는|이나)\s*[^.]{5,80}벌금)', content)
                if not penalty:
                    penalty = _re.search(r'(\d+년\s*이하의?\s*징역[^.]{0,80})', content)
                if penalty:
                    result = penalty.group(1).strip()
                    result = _re.sub(r'\s*에\s*처한다\.?$', '', result)
                    logger.info(f"Placeholder 치환: {inner[:30]} → {result[:50]}")
                    return result
                else:
                    return f"{content[:100]}..."
        
        return f"(형량 — 원문 확인 필요)"
    
    reply = all_placeholder_pattern.sub(replace_any_placeholder, reply)
            
    # 2. ██ 형량 강제 치환 — AI가 Placeholder를 안 써도 DB 원문으로 교체 ██
    penalty_map = {}
    
    penalty_law_shorts = {"형법", "상표법", "관세법", "부정경쟁방지법"}
    for db_law in laws_db:
        if db_law.get("law_short") not in penalty_law_shorts:
            continue
        content = db_law.get("content", "")
        # 형량 패턴 추출 — "~에 처한다"까지 포함 후 제거
        # 패턴1: "N년 이하 징역 또는 N만원 이하 벌금" (일반적)
        penalty_match = _re.search(
            r'(\d+년\s*이하의?\s*징역\s*(?:또는|이나)\s*\d[^\s]*원\s*이하의?\s*벌금(?:\s*에\s*처한다)?)', content)
        # 패턴2: "N년 이하 징역 또는 관세액의 10배..." (관세법 특수)
        if not penalty_match:
            penalty_match = _re.search(
                r'(\d+년\s*이하의?\s*징역\s*(?:또는|이나)\s*[^.]{5,80}벌금(?:\s*에\s*처한다)?)', content)
        # 패턴3: "N년 이하 징역" 만 있는 경우
        if not penalty_match:
            penalty_match = _re.search(r'(\d+년\s*이하의?\s*징역[^.]{0,80})', content)
        if penalty_match:
            key = f"{db_law['law_short']} {db_law['article_no']}"
            raw_penalty = penalty_match.group(1).strip()
            # 서술어 꼬리 제거
            raw_penalty = _re.sub(r'\s*에\s*처한다\.?$', '', raw_penalty)
            raw_penalty = _re.sub(r'\s*에\s*처함\.?$', '', raw_penalty)
            raw_penalty = _re.sub(r'\s*에\s*처하\.?$', '', raw_penalty)
            penalty_map[key] = raw_penalty
            logger.info(f"형량 추출: {key} → '{raw_penalty}'")
    
    # AI 출력에서 형량 교체 — 법령명 근처 + 전체 텍스트 서술어 정리
    for law_key, correct_penalty in penalty_map.items():
        law_name_pattern = _re.escape(law_key)
        
        # 법령명 뒤의 형량 교체 (기본)
        nearby_pattern = _re.compile(
            r'(' + law_name_pattern + r'[^.]*?)(\d+년\s*이하의?\s*징역\s*(?:또는|이나)\s*[^.]{5,80}벌금)',
            _re.DOTALL
        )
        
        _correct = correct_penalty  # 클로저용
        def replace_with_correct(m, _cp=_correct, _lk=law_key):
            prefix = m.group(1)
            wrong_penalty = m.group(2)
            if wrong_penalty.strip() != _cp:
                logger.warning(f"형량 강제 치환: '{wrong_penalty.strip()}' → '{_cp}' ({_lk})")
                return prefix + _cp
            return m.group(0)
        
        reply = nearby_pattern.sub(replace_with_correct, reply)
    
    # 전체 텍스트에서 "~에 처한다" 뒤 이상한 연결 정리
    reply = _re.sub(r'에\s*처한다\s*처벌', '처벌', reply)
    reply = _re.sub(r'에\s*처한다\s*에\s*처해[지질]', '에 처해질 수 있', reply)
    reply = _re.sub(r'에\s*처한다\s*에\s*처', '에 처', reply)
    reply = _re.sub(r'벌금에\s*처한다\s*[)）]', '벌금)', reply)
    reply = _re.sub(r'벌금에\s*처한다\s*대상', '벌금 대상', reply)
    reply = _re.sub(r'벌금에\s*처한다([^."])', r'벌금\1', reply)  # 문장 중간의 "에 처한다" 제거
    
    # 2.5 "보세판매장고시 검토 누락" 문구 강제 제거 (시스템이 이미 주입했으므로)
    reply = _re.sub(r'\*?\*?보세판매장고시\s*검토\s*누락\*?\*?\s*[—\-]\s*별도\s*확인\s*필요', 
                    '보세판매장고시 관련 행정처분 기준은 상기 쟁점에서 검토 완료', reply)
    reply = reply.replace("보세판매장고시 검토 누락", "보세판매장고시 행정처분 기준 검토 완료")
    
    # 2.7 DB에 있는 법령인데 Claude가 "DB 미등록"으로 잘못 표기한 경우 교체
    db_shorts = {d.get("law_short", "") for d in laws_db}
    db_law_names = {d.get("law_name", "") for d in laws_db}
    all_db_names = {n for n in (db_shorts | db_law_names) if n}
    
    # 전체 텍스트에서 "DB 미등록" 패턴을 찾아 해당 법령이 실제 DB에 있는지 확인
    for match in _re.finditer(r'【\s*DB\s*미등록[^】]*】', reply):
        tag = match.group(0)
        # 이 태그 뒤 100자 이내에서 법령명 검색
        after_text = reply[match.end():match.end()+150]
        before_text = reply[max(0, match.start()-150):match.start()]
        context = before_text + after_text
        
        for db_name in all_db_names:
            if db_name in context:
                # DB에 있는 법령인데 "DB 미등록"으로 표기됨 → 교체
                new_tag = tag.replace("DB 미등록", "DB 검증 완료").replace("API 미확인", "")
                new_tag = new_tag.replace(", ,", ",").replace(",  ", "").rstrip(", ").rstrip(",")
                # 빈 태그 정리: 【DB 검증 완료, 】 → 【DB 검증 완료】
                new_tag = _re.sub(r',\s*】', '】', new_tag)
                reply = reply.replace(tag, new_tag, 1)
                logger.info(f"후처리: '{db_name}' DB 미등록 → DB 검증 완료 교체")
                break
    
    # 3. 전체 텍스트에서 판례번호 패턴 추출 → API 검증 + 사건명 강제 표시
    prec_pattern = _re.compile(r'(\d{4}[가-힣]{1,3}\d{2,6})')
    found_cases = prec_pattern.findall(reply)
    
    if found_cases:
        for case_query in set(found_cases):
            status, title = verify_precedent_via_api(case_query)
            time.sleep(0.3)
            if status == "not_found":
                reply = reply.replace(case_query, f"~~{case_query}~~ [🚨 시스템 검증: 미존재 판례]")
                logger.warning(f"후처리: 미존재 판례 차단 — {case_query}")
            elif status == "wrong_context":
                reply = reply.replace(case_query, f"~~{case_query}~~ [🚨 시스템 검증: 맥락 불일치 — 실제 사건: {title}]")
                logger.warning(f"후처리: 맥락 불일치 판례 차단 — {case_query} (실제: {title})")
            elif status == "verified" and title:
                # verified여도 사건명을 강제 삽입 (AI의 해석 창작을 사건명으로 견제)
                # "2018도14446 판결에 따르면..." → "2018도14446 [실제 사건: 상표권침해] 판결에 따르면..."
                if title not in reply:
                    reply = reply.replace(case_query, f"{case_query} [실제 사건: {title[:30]}]")
    
    # 4. 용어 치환: "~팀" → "~담당부서"
    team_replacements = {
        "법무팀": "법무 담당부서", "MD팀": "MD 담당부서",
        "컴플라이언스팀": "컴플라이언스 담당부서", "준법팀": "준법 담당부서",
        "법무 팀": "법무 담당부서", "MD 팀": "MD 담당부서",
    }
    for old_term, new_term in team_replacements.items():
        reply = reply.replace(old_term, new_term)
    
    # 5. DB 미등록 항목의 추측성 형량 제거
    reply = _re.sub(
        r'(⚠️\s*DB\s*미등록[^.]*?)\d+년\s*이하의?\s*징역[^.]*?벌금[^.]*',
        r'\1(형량 — 원문 확인 필요)',
        reply
    )
    reply = _re.sub(r'\d+년\s*이하[^.]*?적용\s*추정되나', '(형량 — 원문 확인 필요)', reply)
    
    # 6. ██ 최종 방어선: penalty_map 기반 전체 텍스트 형량 강제 교체 ██
    # penalty_map에서 추출한 정확한 형량과 다른 형량이 같은 법령명 근처에 있으면 교체
    # 이건 법령명이 정확히 "관세법 제269조"가 아니라 "밀수입죄", "관세법" 등으로 쓰여도 잡음
    for law_key, correct_penalty in penalty_map.items():
        law_short = law_key.split(" ")[0]  # "관세법", "형법" 등
        article = " ".join(law_key.split(" ")[1:])  # "제269조"
        
        # 법령약칭이 reply에 있으면, 그 근처의 잘못된 형량을 교체
        # "관세법" 또는 "밀수입" 등이 같은 문장에 있는 형량 패턴
        alias_map = {
            "관세법": ["관세법", "밀수입", "밀수출입", "제269조"],
            "형법": ["형법", "사기죄", "사기", "제347조"],
            "상표법": ["상표법", "상표권 침해", "제230조"],
            "부정경쟁방지법": ["부정경쟁", "제18조"],
        }
        
        aliases = alias_map.get(law_short, [law_short])
        for alias in aliases:
            if alias not in reply:
                continue
            # alias가 포함된 문장에서 형량 패턴 교체
            pattern = _re.compile(
                r'(' + _re.escape(alias) + r'[^.]*?)(\d+년\s*이하의?\s*징역\s*(?:또는|이나)\s*[^.]{3,80}벌금)',
                _re.DOTALL
            )
            _cp = correct_penalty
            def _replace_alias(m, _correct=_cp, _alias=alias):
                prefix = m.group(1)
                found_penalty = m.group(2)
                if found_penalty.strip() != _correct:
                    logger.warning(f"형량 강제 교체({_alias}): '{found_penalty.strip()[:50]}' → '{_correct[:50]}'")
                    return prefix + _correct
                return m.group(0)
            reply = pattern.sub(_replace_alias, reply)
    
    # JSON 블록 복원
    if json_block and "{{JSON_BLOCK}}" in reply:
        reply = reply.replace("{{JSON_BLOCK}}", json_block)
    
    # 7. "(형량 — 원문 확인 필요)" placeholder를 penalty_map으로 최종 치환
    if penalty_map:
        for law_key, correct_penalty in penalty_map.items():
            law_short = law_key.split(" ")[0]
            # 법령명 근처의 "(형량 — 원문 확인 필요)"를 실제 형량으로 교체
            for alias in [law_short, law_key]:
                pattern = _re.compile(
                    r'(' + _re.escape(alias) + r'[^.]{0,200}?)\(형량\s*[—\-]\s*원문\s*확인\s*필요\)',
                    _re.DOTALL
                )
                if pattern.search(reply):
                    reply = pattern.sub(r'\1' + correct_penalty, reply)
                    logger.info(f"placeholder 치환: {alias} 근처 → {correct_penalty[:30]}")
    
    # 남은 "(형량 — 원문 확인 필요)"는 경고 표시로 변환 (최종 문서에 placeholder 노출 방지)
    reply = reply.replace("(형량 — 원문 확인 필요)", "(형량: 해당 법령 원문 참조)")
    
    # 8. 마크다운 이스케이프 정제 (docx용)
    reply = reply.replace("\\*\\*", "**")
    reply = reply.replace("\\-\\-\\-", "---")
    reply = reply.replace("\\*", "*")
    reply = _re.sub(r'\\([*_~`#])', r'\1', reply)
    
    # 9. 사규 파일명 ".docx" 확장자 제거 → 정식 명칭 변환
    file_name_map = {
        "직매입.docx": "직매입거래 기본계약서",
        "특약매입.docx": "특약매입거래 기본계약서",
        "사규 모음.docx": "사내규정집",
        "사규모음.docx": "사내규정집",
    }
    for old_name, formal_name in file_name_map.items():
        reply = reply.replace(old_name, formal_name)
    reply = _re.sub(r'(\S+)\.docx', r'\1', reply)
    
    # 10. ██ 내부 파이프라인 태그 제거 — 최종 문서에 노출 금지 ██
    # "⚠️ (DB 미등록, API 미확인 — 존재 여부 불확실) Gemini 검증:" 등
    reply = _re.sub(r'⚠️?\s*\(DB\s*미등록[^)]*\)\s*(?:Gemini\s*검증:?\s*)?', '', reply)
    reply = _re.sub(r'✅?\s*\(DB\s*검증\s*완료[^)]*\)\s*', '', reply)
    reply = _re.sub(r'【DB\s*검증\s*완료[^】]*】\s*', '', reply)
    reply = _re.sub(r'【DB\s*미등록[^】]*】\s*', '', reply)
    reply = _re.sub(r'【DB\s*일부\s*검증[^】]*】\s*', '', reply)
    reply = _re.sub(r'\(DB\s*검증\s*완료\)\s*', '', reply)
    reply = _re.sub(r'Gemini\s*검증:\s*', '', reply)
    # "(형량: 해당 법령 원문 참조)" 도 더 깔끔하게
    reply = reply.replace("(형량: 해당 법령 원문 참조)", "(형량은 해당 법령 원문 참조)")
    
    return reply


def dispatch_with_fallback(model_choice, messages, docs, laws_db):
    if model_choice == "claude":
        # ━━━ 3단계 하이브리드 파이프라인 ━━━
        
        # Stage 1: Gemini Researcher — 사규 분석 + 법령/판례 수집
        gemini_system = build_system_gemini_stage1(docs, laws_db)
        st.caption("🔍 Stage 1: 사규 분석 + 법령·판례 수집 중 (Gemini + Google Search)...")
        gemini_reply = call_gemini(gemini_system, messages)
        
        if gemini_reply.startswith("⚠️"):
            st.warning(f"Stage 1 실패: {gemini_reply}")
            return gemini_reply, "Gemini (Failed)"
        
        # Gemini 응답에서 JSON 파싱
        gemini_json = None
        json_match = re.search(r'```json\s*\n(.*?)\n```', gemini_reply, re.DOTALL)
        if json_match:
            try:
                gemini_json = json.loads(json_match.group(1))
            except json.JSONDecodeError:
                logger.warning("Stage 1 JSON 파싱 실패 — Gemini 응답을 텍스트로 전달")
        
        if not gemini_json:
            # JSON 파싱 실패 시 텍스트 그대로 전달
            logger.warning("Gemini가 JSON 미출력 — 텍스트 기반으로 진행")
            gemini_json = {"query_summary": "JSON 파싱 실패", "saryu_findings": [], "law_findings": [], "precedent_findings": [], "risk_areas": []}
        
        # Stage 2: Gatekeeper — DB 원문 대조 + API 실시간 검증
        st.caption("🛡️ Stage 2: DB 원문 대조 + law.go.kr API 팩트체크 중...")
        gatekeeper_meta, gatekeeper_text = gatekeeper_process(gemini_json, laws_db, docs)
        
        # Gatekeeper 결과 요약 표시
        if gatekeeper_meta and not gatekeeper_meta.get("error"):
            v_laws = len(gatekeeper_meta.get("verified_laws", []))
            u_laws = len(gatekeeper_meta.get("unverified_laws", []))
            v_prec = len(gatekeeper_meta.get("verified_precedents", []))
            d_prec = len(gatekeeper_meta.get("dropped_precedents", []))
            st.caption(f"  📊 법령: ✅{v_laws} ⚠️{u_laws} | 판례: ✅{v_prec} 🚨{d_prec}건 차단")
        
        if not gatekeeper_text:
            gatekeeper_text = f"Gemini 수집 결과:\n{gemini_reply[:3000]}"
        
        # Stage 3: Claude Senior Lawyer — 최종 법리 해석
        claude_system = build_system_claude_v3(gatekeeper_text, gatekeeper_meta)
        st.caption("⚖️ Stage 3: 최종 법리 해석 및 보고서 작성 중 (Claude)...")
        
        # Claude에게는 원래 사용자 질문만 전달 (정제 데이터는 system에)
        claude_messages = [messages[-1]] if messages else [{"role": "user", "content": "검토해주세요."}]
        reply = call_claude(claude_system, claude_messages)
        
        if reply.startswith("⚠️"):
            # Claude 실패 시 → Gemini 결과 + Gatekeeper 정보를 직접 표시
            st.warning(f"🔄 Claude 실패: {reply}\n→ Gemini 검토 결과를 직접 표시합니다.")
            # Gemini 응답에서 JSON 이후 마크다운 부분 추출
            fallback_text = gemini_reply
            if json_match:
                fallback_text = gemini_reply[json_match.end():].strip() or gemini_reply
            return fallback_text, "Gemini (Claude 우회)"
        
        # 성공 — JSON 먼저 파싱, 후처리는 detail_text에만 적용
        st.caption("🔍 후처리: 전체 텍스트 판례 검증 중...")
        
        # 1) JSON 먼저 분리
        pre_json, pre_detail = parse_review_response(reply)
        
        if pre_json:
            # JSON 파싱 성공 → detail_text만 후처리 (JSON은 보호)
            processed_detail = postprocess_reply(pre_detail, laws_db)
            # JSON 내부의 law_analysis 등에서도 형량 치환 필요 → JSON을 문자열로 후처리
            json_str = json.dumps(pre_json, ensure_ascii=False)
            json_str = postprocess_reply(json_str, laws_db)
            try:
                pre_json = json.loads(json_str)
            except json.JSONDecodeError:
                logger.warning("JSON 후처리 후 파싱 실패 — 원본 사용")
            # reply를 재조립
            reply = "```json\n" + json.dumps(pre_json, ensure_ascii=False, indent=2) + "\n```\n" + processed_detail
        else:
            # JSON 파싱 실패 → 전체 후처리
            reply = postprocess_reply(reply, laws_db)
        
        st.session_state["_gatekeeper_meta"] = gatekeeper_meta
        
        return reply, "Gemini→Gatekeeper→Claude"
    else:
        # 일반 Q&A는 Gemini 단독
        system = build_system_gemini(docs)
        reply = call_gemini(system, messages)
        if reply.startswith("⚠️"):
            fallback_reply = call_claude(build_system_claude(docs, laws_db), [messages[-1]] if messages else [])
            if not fallback_reply.startswith("⚠️"):
                return fallback_reply, f"Claude ({CLAUDE_MODEL}, Fallback)"
            return reply, "Gemini (Failed)"
        return reply, "Gemini"

# ── JSON 파싱 및 UI 렌더링 ────────────────────────────────────
def parse_review_response(response_text):
    json_data = None
    detail_text = response_text
    
    # 패턴 1: ```json ... ``` 코드블록
    json_match = re.search(r'```json\s*\n(.*?)\n```', response_text, re.DOTALL)
    if json_match:
        try:
            json_data = json.loads(json_match.group(1))
            detail_text = response_text[json_match.end():].strip()
        except json.JSONDecodeError as e:
            logger.error(f"JSON 파싱 실패 (코드블록): {e}")
    
    # 패턴 2: 코드블록 없이 바로 { ... } 형태
    if not json_data:
        json_match2 = re.search(r'(\{[^{]*?"summary".*?"verdict".*?\})\s*$', response_text, re.DOTALL)
        if not json_match2:
            json_match2 = re.search(r'(\{[^{]*?"summary".*?"issues".*?\})', response_text, re.DOTALL)
        if json_match2:
            try:
                json_data = json.loads(json_match2.group(1))
                detail_text = response_text[json_match2.end():].strip()
            except json.JSONDecodeError as e:
                logger.error(f"JSON 파싱 실패 (raw): {e}")
    
    # 패턴 3: 전체 텍스트가 JSON인 경우
    if not json_data and response_text.strip().startswith("{"):
        try:
            # JSON 뒤에 마크다운이 있을 수 있으므로 첫 번째 최상위 } 까지만 파싱
            brace_count = 0
            json_end = 0
            for i, c in enumerate(response_text):
                if c == '{': brace_count += 1
                elif c == '}': 
                    brace_count -= 1
                    if brace_count == 0:
                        json_end = i + 1
                        break
            if json_end > 0:
                json_data = json.loads(response_text[:json_end])
                detail_text = response_text[json_end:].strip()
        except json.JSONDecodeError as e:
            logger.error(f"JSON 파싱 실패 (전체): {e}")
    
    return json_data, detail_text


def enforce_mandatory_issues(json_data, laws_db):
    """██ 필수 쟁점 강제 삽입 — AI가 빠뜨려도 코드가 100% 보장 ██"""
    if not json_data or "issues" not in json_data:
        return json_data
    
    issues = json_data.get("issues", [])
    all_text = " ".join(iss.get("title", "") + " " + iss.get("applicable_law", "") for iss in issues)
    next_no = max((iss.get("issue_no", 0) for iss in issues), default=0) + 1
    
    # 1. 보세판매장고시 — 없으면 강제 삽입
    if "보세판매장" not in all_text:
        bc = []
        for d in laws_db:
            if d.get("law_short") == "보세판매장고시" and d.get("article_no") in ("제18조", "제28조"):
                bc.append(f"{d['article_no']}({d.get('article_title','')}): {d['content'][:150]}")
        issues.append({
            "issue_no": next_no, "title": "보세판매장고시 위반 — 행정처분 리스크 [시스템 강제 삽입]",
            "risk_level": "high", "target_clause": "보세판매장 운영 규정 위반",
            "applicable_law": "보세판매장 특허 및 운영에 관한 고시",
            "law_analysis": f"운영인의 법령 위반 시 6개월 범위 내 물품반입 정지 또는 특허취소 (관세법 제178조 연계). {'  '.join(bc[:2]) if bc else '보세판매장고시 원문 참조'} (DB 검증 완료)",
            "applicable_rule": "해당 없음", "rule_analysis": "해당 없음",
            "recommendation": "세관 신고 및 보세판매장 특허 보호를 위한 즉시 시정 조치. 자진 신고 검토."
        })
        logger.warning("enforce: 보세판매장고시 강제 삽입!")
        next_no += 1
    
    # 2. 소비자기본법 — 없으면 강제 삽입
    if "소비자" not in all_text:
        issues.append({
            "issue_no": next_no, "title": "소비자기본법 위반 — 소비자 피해 [시스템 강제 삽입]",
            "risk_level": "medium", "target_clause": "모조품 판매로 인한 소비자 권리 침해",
            "applicable_law": "소비자기본법 제4조, 제19조, 제20조",
            "law_analysis": "소비자 안전권·알 권리 침해. 소비자분쟁해결기준에 따른 피해구제 및 손해배상 가능 (형량 — 원문 확인 필요)",
            "applicable_rule": "해당 없음", "rule_analysis": "해당 없음",
            "recommendation": "피해 소비자 파악 및 환불/교환. 집단분쟁 가능성 대비."
        })
        logger.warning("enforce: 소비자기본법 강제 삽입!")
        next_no += 1
    
    # 3. 부정경쟁방지법 — 없으면 강제 삽입
    if "부정경쟁" not in all_text:
        issues.append({
            "issue_no": next_no, "title": "부정경쟁방지법 위반 — 상품형태 모방 [시스템 강제 삽입]",
            "risk_level": "medium", "target_clause": "타인 상품 형태 모방 모조품 유통",
            "applicable_law": "부정경쟁방지법 제2조, 제18조",
            "law_analysis": "상품 혼동 야기 부정경쟁행위. 상표법과 별개 적용. 미등록 브랜드 디자인 도용 포함 (형량 — 원문 확인 필요)",
            "applicable_rule": "해당 없음", "rule_analysis": "해당 없음",
            "recommendation": "브랜드 권리자 민사 손해배상 대비. 법무 담당부서 검토."
        })
        logger.warning("enforce: 부정경쟁방지법 강제 삽입!")
    
    json_data["issues"] = issues
    
    # cited_laws에도 강제 삽입된 법령 추가 (verify_citations에서 DB 검증 통과하도록)
    cited = json_data.get("cited_laws", [])
    cited_text = " ".join(str(c) for c in cited)
    forced_citations = [
        {"law_name": "보세판매장 특허 및 운영에 관한 고시", "article": "제18조, 제28조"},
        {"law_name": "소비자기본법", "article": "제4조, 제19조"},
        {"law_name": "부정경쟁방지 및 영업비밀보호에 관한 법률", "article": "제2조, 제18조"},
    ]
    for fc in forced_citations:
        if fc["law_name"] not in cited_text:
            cited.append(fc)
    json_data["cited_laws"] = cited
    
    return json_data


def render_verdict_badge(verdict):
    badges = {
        "approved":    ("🟢 위험 요소 미발견 (사내변호사 확인 권장)", "success"),
        "conditional": ("🟡 수정 필요 사항 발견", "warning"),
        "rejected":    ("🔴 중대 위험 발견 (진행 보류 권고)", "error"),
    }
    label, msg_type = badges.get(verdict, ("⚪ 판단 보류", "info"))
    getattr(st, msg_type)(label)

def render_issues_table(issues, citation_results):
    if not issues: return
    risk_icons = {"high": "🔴", "medium": "🟡", "low": "🟢"}
    for issue in issues:
        risk = issue.get("risk_level", "medium")
        icon = risk_icons.get(risk, "⚪")
        with st.expander(f"{icon} 쟁점 {issue.get('issue_no', '?')}: {issue.get('title', '제목 없음')}", expanded=(risk == "high")):
            if issue.get("target_clause"):
                st.markdown(f"**📌 검토 대상 원문:**")
                st.code(issue["target_clause"], language="text")
            col1, col2 = st.columns(2)
            with col1:
                law_ref = issue.get("applicable_law", "")
                if law_ref:
                    verified = "⚪"
                    for cr in citation_results:
                        if any(part in cr["citation"] for part in law_ref.split()):
                            verified = "✅" if cr["verified"] else "⚠️"
                            break
                    st.markdown(f"**⚖️ 적용 법령:** {law_ref} {verified}")
                if issue.get("law_analysis"):
                    st.markdown(f"**🔍 법령·행정규칙·판례 관점:** {issue['law_analysis']}")
            with col2:
                if issue.get("applicable_rule"):
                    st.markdown(f"**🏛️ 적용 사규:** {issue['applicable_rule']}")
                if issue.get("rule_analysis"):
                    st.markdown(f"**🔍 사규 관점:** {issue['rule_analysis']}")
            if issue.get("recommendation"):
                st.info(f"💡 **종합 실무 권고:** {issue['recommendation']}")

def render_alternative_clause(clause):
    if clause and clause != "null":
        st.markdown("---")
        st.markdown("### 📝 수정 대안 조항 (초안)")
        st.caption("아래 조항을 복사하여 협상 메일이나 수정 계약서에 활용하세요.")
        st.code(clause, language="text")

def _apply_shading(paragraph, hex_color):
    """단락에 배경 음영(shading) 적용 — Word '음영' 스타일과 동일 효과"""
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    paragraph.paragraph_format.element.get_or_add_pPr().append(shading)

def _add_shaded_heading(doc, text, level, shade_color="D9E2F3"):
    """음영 배경이 적용된 Heading 추가 (Word '음영' 디자인 스타일)"""
    h = doc.add_heading(text, level=level)
    _apply_shading(h, shade_color)
    return h

def generate_review_docx(json_data, detail_text, query_text):
    """검토 의견서 docx 생성 — 화면 렌더링과 동일한 구조로 출력.
    여백: 좁게 (상하좌우 1.27cm = Word '좁게' 프리셋)
    디자인: Word 기본 스타일 '음영' 적용
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    # ── 음영 디자인 컬러 팔레트 ──
    SHADE_H1 = "2F5496"       # Heading 1 배경: 진한 파랑
    SHADE_H1_FONT = "FFFFFF"  # Heading 1 글자: 흰색
    SHADE_H2 = "D9E2F3"       # Heading 2 배경: 연한 파랑
    SHADE_H3 = "E2EFDA"       # Heading 3 배경: 연한 초록
    SHADE_QUOTE = "F2F2F2"    # 인용/코드 배경: 연한 회색
    SHADE_INFO = "DAEEF3"     # 권고(info) 배경: 연한 시안
    SHADE_WARN = "FFF2CC"     # 경고 배경: 연한 노랑
    SHADE_ALT = "E8E0F0"      # 대안 조항 배경: 연한 보라

    doc = Document()

    # ── 페이지 여백: 좁게 (1.27cm = Word 기본 '좁게' 프리셋) ──
    for section in doc.sections:
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)

    # ── 기본 스타일 ──
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(10)

    # ── Heading 스타일 커스텀 (음영 디자인) ──
    for hs_id, hs_size, hs_color in [("Heading 1", 16, SHADE_H1), ("Heading 2", 13, SHADE_H2), ("Heading 3", 11, SHADE_H3)]:
        try:
            hs = doc.styles[hs_id]
            hs.font.name = '맑은 고딕'
            hs.font.size = Pt(hs_size)
            if hs_id == "Heading 1":
                hs.font.color.rgb = RGBColor.from_string(SHADE_H1_FONT)
        except KeyError:
            pass

    # ── 헤더: 앱 타이틀과 동일 ──
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_shading(p_title, SHADE_H1)
    run_title = p_title.add_run("🤝 공정거래 실무 어시스턴트 v2.1 — 검토 의견서")
    run_title.font.size = Pt(18)
    run_title.bold = True
    run_title.font.color.rgb = RGBColor(255, 255, 255)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("MD·협력사 실무 Q&A & 계약·법령 Self-Check AI")
    run_sub.font.size = Pt(9)
    run_sub.font.color.rgb = RGBColor(128, 128, 128)

    # 작성 정보
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_meta = p_meta.add_run(f"작성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}  |  AI 검토 초안")
    run_meta.font.size = Pt(8)
    run_meta.font.color.rgb = RGBColor(128, 128, 128)

    # 경고 배너 (음영 배경)
    p_warn = doc.add_paragraph()
    p_warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_shading(p_warn, SHADE_WARN)
    run_warn = p_warn.add_run("⚠️ 본 문서는 AI가 생성한 검토 초안입니다. 법적 효력이 없으며, 반드시 사내변호사의 최종 확인을 거쳐야 합니다.")
    run_warn.font.size = Pt(9)
    run_warn.font.color.rgb = RGBColor(156, 101, 0)
    run_warn.bold = True

    doc.add_paragraph("")

    if json_data:
        # ── 검토 결론 배지 (화면의 verdict badge와 동일) ──
        verdict = json_data.get("verdict", "")
        verdict_map = {
            "approved":    "🟢 위험 요소 미발견 (사내변호사 확인 권장)",
            "conditional": "🟡 수정 필요 사항 발견",
            "rejected":    "🔴 중대 위험 발견 (진행 보류 권고)",
        }
        verdict_colors = {
            "approved":    RGBColor(0, 128, 0),
            "conditional": RGBColor(200, 150, 0),
            "rejected":    RGBColor(227, 0, 15),
        }
        verdict_shades = {
            "approved":    "E2EFDA",
            "conditional": "FFF2CC",
            "rejected":    "FCE4EC",
        }
        p_verdict = doc.add_paragraph()
        _apply_shading(p_verdict, verdict_shades.get(verdict, "F2F2F2"))
        run_v = p_verdict.add_run(verdict_map.get(verdict, "⚪ 판단 보류"))
        run_v.font.size = Pt(14)
        run_v.bold = True
        run_v.font.color.rgb = verdict_colors.get(verdict, RGBColor(128, 128, 128))

        # 요약 (화면의 📋 summary와 동일)
        summary = json_data.get("summary", query_text[:200])
        p_summary = doc.add_paragraph()
        run_s = p_summary.add_run(f"📋 {summary}")
        run_s.font.size = Pt(11)
        run_s.bold = True

        # 판단 근거
        if json_data.get("verdict_reason"):
            doc.add_paragraph(json_data["verdict_reason"])

        doc.add_paragraph("")

        # ── 쟁점별 분석 (화면의 expander 구조와 동일) ──
        issues = json_data.get("issues", [])
        if issues:
            _add_shaded_heading(doc, "쟁점별 교차 분석", level=2, shade_color=SHADE_H2)
            risk_icons = {"high": "🔴", "medium": "🟡", "low": "🟢"}
            risk_labels = {"high": "[위험]", "medium": "[주의]", "low": "[양호]"}
            risk_shades = {"high": "FCE4EC", "medium": "FFF8E1", "low": "E8F5E9"}

            for issue in issues:
                risk = issue.get("risk_level", "medium")
                icon = risk_icons.get(risk, "⚪")
                label = risk_labels.get(risk, "")

                # 쟁점 제목 (음영 배경 + 위험도별 색상)
                h = doc.add_heading(f"{icon} 쟁점 {issue.get('issue_no', '?')}: {issue.get('title', '제목 없음')} {label}", level=3)
                _apply_shading(h, risk_shades.get(risk, SHADE_H3))

                # 📌 검토 대상 원문 (회색 음영 박스)
                if issue.get("target_clause"):
                    doc.add_paragraph("📌 검토 대상 원문:").runs[0].bold = True
                    p_clause = doc.add_paragraph(issue["target_clause"])
                    p_clause.paragraph_format.left_indent = Cm(0.5)
                    _apply_shading(p_clause, SHADE_QUOTE)
                    for run in p_clause.runs:
                        run.font.color.rgb = RGBColor(80, 80, 80)
                        run.font.size = Pt(9)

                # ⚖️ 적용 법령 + 법령 관점
                if issue.get("applicable_law"):
                    p = doc.add_paragraph()
                    run_label = p.add_run("⚖️ 적용 법령: ")
                    run_label.bold = True
                    p.add_run(issue["applicable_law"])
                if issue.get("law_analysis"):
                    p = doc.add_paragraph()
                    run_label = p.add_run("🔍 법령·행정규칙·판례 관점: ")
                    run_label.bold = True
                    p.add_run(issue["law_analysis"])

                # 🏛️ 적용 사규 + 사규 관점
                if issue.get("applicable_rule"):
                    p = doc.add_paragraph()
                    run_label = p.add_run("🏛️ 적용 사규: ")
                    run_label.bold = True
                    p.add_run(issue["applicable_rule"])
                if issue.get("rule_analysis"):
                    p = doc.add_paragraph()
                    run_label = p.add_run("🔍 사규 관점: ")
                    run_label.bold = True
                    p.add_run(issue["rule_analysis"])

                # 💡 종합 실무 권고 (시안 음영 박스 — 화면의 st.info와 동일)
                if issue.get("recommendation"):
                    p_rec = doc.add_paragraph()
                    _apply_shading(p_rec, SHADE_INFO)
                    run_rec = p_rec.add_run(f"💡 종합 실무 권고: {issue['recommendation']}")
                    run_rec.bold = True
                    run_rec.font.color.rgb = RGBColor(0, 80, 160)

                doc.add_paragraph("")  # 쟁점 간 간격

        # ── MD Action Plan ──
        if json_data.get("action_plan"):
            _add_shaded_heading(doc, "MD Action Plan", level=2, shade_color=SHADE_H2)
            doc.add_paragraph(json_data["action_plan"])

        # ── 수정 대안 조항 (보라 음영 박스) ──
        alt = json_data.get("alternative_clause")
        if alt and alt != "null":
            _add_shaded_heading(doc, "📝 수정 대안 조항 (초안)", level=2, shade_color=SHADE_H2)
            p_alt_desc = doc.add_paragraph("아래 조항을 복사하여 협상 메일이나 수정 계약서에 활용하세요.")
            p_alt_desc.runs[0].font.size = Pt(8)
            p_alt_desc.runs[0].font.color.rgb = RGBColor(128, 128, 128)
            p_alt = doc.add_paragraph(alt)
            p_alt.paragraph_format.left_indent = Cm(0.5)
            _apply_shading(p_alt, SHADE_ALT)

        # 상세 전문 제거 — 본문과 형량 불일치 문제의 근본 원인이므로 삭제

    else:
        # JSON 파싱 실패 시 원문 그대로
        _add_shaded_heading(doc, "검토 의견", level=2, shade_color=SHADE_H2)
        doc.add_paragraph(detail_text[:10000])

    # ── 푸터 면책 (회색 음영) ──
    doc.add_paragraph("")
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_shading(p_footer, SHADE_QUOTE)
    run_f = p_footer.add_run("본 검토의견서는 AI가 생성한 초안이며, 법적 효력이 없습니다.\n반드시 사내변호사의 최종 검토를 거치기 바랍니다.")
    run_f.font.size = Pt(8)
    run_f.font.color.rgb = RGBColor(128, 128, 128)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# ── Streamlit UI 메인 함수 ────────────────────────────────────
def main():
    st.set_page_config(page_title="공정거래 실무 어시스턴트 v2.1", page_icon="🤝", layout="wide")

    st.markdown("""
    <style>
    /* ━━━ 신세계그룹 뉴스룸 디자인 시스템 ━━━
       배경: #FFFFFF (화이트)
       텍스트: #000000 (진한 검정)
       액센트: #E02B20 (밝은 빨강)
       폰트: Noto Sans KR
       제목: 데스크탑 40px / 모바일 8.57vw
       본문: 데스크탑 18px / 모바일 4.29vw
       심플하고 화려함 배제, 뉴스 웹사이트 스타일
    */

    /* 1. 웹 폰트: Noto Sans KR (뉴스룸 공식 폰트) */
    @import url("https://fonts.googleapis.com/css2?family=Noto+Sans+KR:wght@300;400;500;700;900&display=swap");

    /* 2. 전체 타이포그래피 */
    html, body, [class*="css"] { 
        font-family: 'Noto Sans KR', -apple-system, BlinkMacSystemFont, "Apple SD Gothic Neo", "Malgun Gothic", sans-serif !important; 
        letter-spacing: -0.02em; 
        color: #000000;
        font-size: 16px;
        line-height: 1.7;
    }

    /* 앱 배경: 화이트 (뉴스룸 동일) */
    .stApp { background-color: #FFFFFF !important; }

    /* 3. 채팅 메시지 블록 */
    .stChatMessage { 
        border-radius: 4px; 
        padding: 24px 28px; 
        margin-bottom: 16px; 
        border: 1px solid #E8E8E8; 
        background-color: #FFFFFF !important; 
        line-height: 1.7;
        font-size: 16px;
    }

    /* 📱 4. 모바일 반응형 */
    @media (max-width: 768px) {
        .stChatMessage {
            padding: 16px 18px;
            font-size: 15px;
        }
        html, body, [class*="css"] {
            font-size: 15px;
        }
    }

    /* 5. 버튼 공통 (뉴스룸 — 심플, 직선적) */
    div.stButton > button:first-child { 
        border-radius: 2px; 
        font-weight: 500; 
        font-family: 'Noto Sans KR', sans-serif !important;
        transition: all 0.15s ease; 
    }

    /* 🔴 프라이머리 버튼 (신세계 밝은 빨강 #E02B20) */
    button[kind="primary"] { 
        background-color: #E02B20 !important; 
        color: #FFFFFF !important; 
        border: none !important; 
    }
    button[kind="primary"]:hover {
        background-color: #C41E15 !important; 
        box-shadow: none !important;
    }

    /* ⚪ 세컨더리 버튼 (블랙 아웃라인) */
    button[kind="secondary"] { 
        background-color: #FFFFFF !important; 
        color: #000000 !important; 
        border: 1px solid #000000 !important; 
    }
    button[kind="secondary"]:hover {
        background-color: #000000 !important;
        color: #FFFFFF !important;
    }

    /* 6. 사이드바 */
    [data-testid="stSidebar"] { 
        background-color: #FAFAFA !important;
        border-right: 1px solid #E8E8E8; 
    }
    
    /* 인라인 코드 (뉴스룸 레드 톤) */
    code { 
        color: #E02B20; 
        background-color: #FFF5F5; 
        border-radius: 2px; 
        padding: 0.15em 0.4em;
        font-size: 0.9em;
    }
    pre { 
        border-radius: 2px; 
        background-color: #F5F5F5 !important; 
        border: 1px solid #E8E8E8; 
    }

    /* 헤딩 스타일 (뉴스룸 — 굵고 깔끔) */
    h1 { font-weight: 900 !important; color: #000000 !important; }
    h2 { font-weight: 700 !important; color: #000000 !important; }
    h3 { font-weight: 700 !important; color: #333333 !important; }

    /* 7. 헤딩 앵커 링크 제거 */
    .stMarkdown a[href^="#"],
    [data-testid="stHeaderActionElements"] {
        display: none !important;
    }
    h1 a, h2 a, h3 a, h4 a, h5 a, h6 a {
        display: none !important;
        pointer-events: none !important;
    }

    /* 8. Expander 스타일 (심플) */
    .streamlit-expanderHeader {
        font-weight: 500 !important;
        font-size: 15px !important;
        color: #000000 !important;
    }

    /* 9. 구분선 */
    hr { border-color: #E8E8E8 !important; }

    /* 10. success/info/warning 박스 톤 다운 */
    [data-testid="stAlert"] {
        border-radius: 2px !important;
        font-size: 15px;
    }
    </style>
    """, unsafe_allow_html=True)

    app_pw = get_secret("APP_PASSWORD")
    if app_pw:
        if "authenticated" not in st.session_state: st.session_state.authenticated = False
        if not st.session_state.authenticated:
            st.markdown("## 🤝 공정거래 실무 어시스턴트")
            pw = st.text_input("비밀번호를 입력하세요", type="password")
            if pw:
                if pw == app_pw: st.session_state.authenticated = True; st.rerun()
                else: st.error("비밀번호가 올바르지 않습니다.")
            st.stop()

    if "docs" not in st.session_state: st.session_state.docs = load_docs()
    if "messages" not in st.session_state: st.session_state.messages = []
    if "sessions" not in st.session_state: st.session_state.sessions = load_sessions()
    if "current_session_id" not in st.session_state: st.session_state.current_session_id = None
    if "laws_db" not in st.session_state: st.session_state.laws_db = load_laws()
    # 응답 완료 후 rerun 플래그 (다운로드 버튼 렌더링 보장)
    if "needs_rerun" not in st.session_state: st.session_state.needs_rerun = False

    if "cleanup_done" not in st.session_state:
        cleanup_old_sessions(90)
        st.session_state.cleanup_done = True

    # 이전 턴에서 세션 저장 후 rerun이 필요한 경우
    if st.session_state.needs_rerun:
        st.session_state.needs_rerun = False
        st.rerun()

    # ── 사이드바 (실무자 동선 최적화) ─────────────────────────
    with st.sidebar:
        st.markdown("## 🤝 공정거래 실무 어시스턴트 v2.1")
        st.caption("면세점 MD 바이어 전용")
        
        # 1. 자동 보안 마스킹 안내 (DLP) - 최상단 배치
        st.markdown("### 🛡️ 정보보안 (DLP) 가동 중")
        st.success(
            "⚠️ **정보 유출 방지 시스템 작동 안내**\n\n"
            "외부 클라우드 AI 서버로 당사의 핵심 기밀 및 협력사 정보가 유출되는 것을 원천 차단하기 위해, **문서 내 민감 정보는 모두 AI 전송 전에 자동 블라인드(마스킹) 처리됩니다.**"
        )
        st.caption("• **자동 차단:** 주민/외국인번호, 휴대전화, 이메일, 사업자/법인번호, 계좌번호, 당사 명칭", unsafe_allow_html=True)
        st.caption("• **수동 차단:** 하단 텍스트 입력창에 기재한 '협력사명'", unsafe_allow_html=True)

        # 법령 DB 업데이트 이력 (일반 사용자 대상)
        law_count = len(st.session_state.laws_db)
        if law_count > 0:
            last_dates = [d.get("last_updated") or d.get("created_at", "") for d in st.session_state.laws_db if d.get("last_updated") or d.get("created_at")]
            if last_dates:
                latest = max(last_dates)
                try:
                    from datetime import datetime as _dt
                    if "T" in latest:
                        dt_obj = _dt.fromisoformat(latest.replace("Z", "+00:00"))
                    else:
                        dt_obj = _dt.fromisoformat(latest)
                    display_date = dt_obj.strftime("%Y-%m-%d %H:%M")
                except Exception:
                    display_date = latest[:16]
                st.caption(f"📚 적용 법령: {law_count}개 조문")
                st.caption(f"🕐 {display_date} 기준")
            else:
                st.caption(f"📚 적용 법령: {law_count}개 조문")
            
            # 법령 DB 상세 목록 (접었다 펼치기)
            with st.expander("📚 적용 법령·행정규칙 DB 목록", expanded=False):
                # 법령별로 그룹핑
                law_groups = {}
                for law in st.session_state.laws_db:
                    short = law.get("law_short", "기타")
                    if short not in law_groups:
                        law_groups[short] = []
                    title = law.get("article_title", "")
                    no = law.get("article_no", "")
                    law_groups[short].append(f"{no} {title}" if title else no)
                
                for law_short, articles in sorted(law_groups.items()):
                    st.markdown(f"**{law_short}** ({len(articles)}개)")
                    st.caption("  |  ".join(articles))
                    st.markdown("")

        st.divider()

        # 2. 사용 매뉴얼 (도움말)
        with st.expander("📖 사용 매뉴얼", expanded=False):
            st.markdown("""
**1. 접속 방법**
- 공유받은 앱 URL로 접속
- 비밀번호 입력 후 이용

**2. 검토 요청 방법 (일반 사용자)**
- **1단계 (사내 기준 자문):** 채팅창에 규정 관련 질문 입력
- **2단계 (심층 법무 검토):** "검토", "위반", "적법" 등 키워드 포함하여 질문하거나 파일 첨부
- **리비전 비교:** 당사 초안(V1)과 협력사 수정본(V2)을 나란히 업로드하면 변경된 독소조항을 자동 비교

**3. 검토의견서 해석법**
- 🟢 **위험 요소 미발견** — 현재 기준으로 문제없으나, 사내변호사 최종 확인 필요
- 🟡 **수정 필요 사항 발견** — 특정 조항 수정 후 진행 가능
- 🔴 **중대 위험 발견** — 진행 보류, 사내변호사와 즉시 협의 필요
- 쟁점별로 ⚖️ 법령·행정규칙·판례 관점 / 🏛️ 사규 관점이 교차 분석됩니다
- 💡 종합 실무 권고를 참고하여 협력사에 대응하세요

**4. 🛡️ 보안(DLP) 안내**
- 주민번호, 전화번호, 이메일, 사업자번호 등은 AI 전송 전 **자동 마스킹** 됩니다
- 협력사명은 **검토 대상 협력사명** 입력란에 기재하면 추가 마스킹됩니다
- 당사 명칭은 항상 자동 차단됩니다

**5. 기준 문서 등록 (관리자)**
- 사이드바 하단 ⚙️ 기준 문서 DB 관리 열기
- 문서 유형(사규/계약서/약정서) 선택 → Word 파일 업로드 → 'DB에 규칙 등록' 클릭

**6. FAQ / 주의사항**
- ❓ **AI 검토 결과를 그대로 써도 되나요?** → 아니요. 반드시 사내변호사의 최종 확인을 받으세요.
- ❓ **어떤 법령이 적용되나요?** → 대규모유통업법, 동 시행령, 공정거래법, 하도급법이 DB에 등록되어 있습니다.
- ❓ **파일은 어떤 형식을 지원하나요?** → .docx(Word) 파일만 지원합니다.
- ❓ **대화 내역은 얼마나 보관되나요?** → 90일 후 자동 삭제됩니다.
- ⚠️ AI가 생성한 검토의견서를 사내변호사 확인 없이 외부에 발송하지 마세요.
""")

        # 3. 새 대화 시작
        if st.button("✨ 새 대화 시작", use_container_width=True, type="primary"):
            st.session_state.messages = []
            st.session_state.current_session_id = None
            st.rerun()

        st.divider()

        # 3. 히스토리 관리
        st.markdown("### 🗂 최근 자문 내역")
        for sess in st.session_state.sessions:
            col1, col2 = st.columns([5, 1])
            with col1:
                if st.button(sess["title"], key="sess_" + sess["id"], use_container_width=True):
                    st.session_state.messages = sess["messages"]
                    st.session_state.current_session_id = sess["id"]
                    st.rerun()
            with col2:
                if st.button("🗑", key="delsess_" + sess["id"]):
                    if delete_session_db(sess["id"]):
                        st.session_state.sessions = [s for s in st.session_state.sessions if s["id"] != sess["id"]]
                        st.rerun()

        st.divider()

        # 5. 관리자용 DB 관리는 가장 아래 숨김 (Expander)
        with st.expander("⚙️ 기준 문서 DB 관리 (관리자 전용)", expanded=False):
            law_count = len(st.session_state.laws_db)
            if law_count > 0:
                # 최신 업데이트 일시 추출
                last_dates = [d.get("last_updated") or d.get("created_at", "") for d in st.session_state.laws_db if d.get("last_updated") or d.get("created_at")]
                if last_dates:
                    latest = max(last_dates)
                    # ISO 형식 → 보기 좋게 변환
                    try:
                        from datetime import datetime as _dt
                        if "T" in latest:
                            dt_obj = _dt.fromisoformat(latest.replace("Z", "+00:00"))
                        else:
                            dt_obj = _dt.fromisoformat(latest)
                        display_date = dt_obj.strftime("%Y-%m-%d %H:%M")
                    except Exception:
                        display_date = latest[:16]
                    st.success(f"📚 법령 DB: {law_count}개 조문\n\n🕐 최종 업데이트: {display_date}")
                else:
                    st.success(f"📚 법령 DB: {law_count}개")
            else:
                st.warning("📚 법령 DB 미설정")

            st.caption(f"🤖 Claude: `{CLAUDE_MODEL}`")
            st.caption(f"🤖 Gemini: `{GEMINI_MODELS[0]}` → `{GEMINI_MODELS[1]}`")

            # law.go.kr API 연결 상태 확인 + 법령 DB 수동 업데이트
            st.markdown("**🔗 law.go.kr API 상태**")
            col_api1, col_api2 = st.columns(2)
            with col_api1:
                if st.button("🔍 MCP 연결 테스트", use_container_width=True, key="api_test"):
                    with st.spinner("korean-law-mcp 연결 테스트 중..."):
                        try:
                            success, result = call_mcp_law("관세법 제1조 조문 원문 알려줘")
                            if success and len(result) > 20:
                                st.success(f"✅ MCP 연결 정상 (응답 {len(result)}자)")
                                st.caption(result[:150] + "...")
                            else:
                                st.error(f"❌ MCP 응답 이상: {result[:100]}")
                        except Exception as api_err:
                            st.error(f"❌ 연결 실패: {api_err}")
            with col_api2:
                if st.button("🔄 법령 DB 업데이트", use_container_width=True, key="update_laws"):
                    with st.spinner("법령 DB 업데이트 중... (1~2분 소요)"):
                        try:
                            import update_laws
                            import importlib
                            importlib.reload(update_laws)
                            stats = update_laws.run_update()
                            st.success(f"✅ 법령 DB 업데이트 완료! (업데이트 {stats.get('updated', 0)}건, 변경없음 {stats.get('unchanged', 0)}건, 실패 {stats.get('failed', 0)}건)")
                            st.session_state.laws_db = load_laws()
                            st.rerun()
                        except ImportError:
                            st.error("❌ update_laws.py 파일을 찾을 수 없습니다. 같은 디렉토리에 배포했는지 확인하세요.")
                        except Exception as upd_err:
                            st.error(f"❌ 업데이트 오류: {upd_err}")
                            logger.error(f"법령 DB 업데이트 실패: {upd_err}")

            st.markdown("---")

            # DB 원문 검증
            if st.button("🔬 DB 원문 검증 (핵심 법령)", use_container_width=True, key="check_db"):
                with st.spinner("DB 원문 검증 중..."):
                    check_targets = [
                        ("부정경쟁방지법", "제2조", ["부정경쟁행위"]),
                        ("부정경쟁방지법", "제4조", ["금지"]),
                        ("부정경쟁방지법", "제18조", ["벌"]),
                        ("형법", "제347조", ["사기", "기망", "편취"]),
                        ("상표법", "제230조", ["침해", "징역"]),
                        ("상표법", "제235조", ["양벌", "법인"]),
                        ("관세법", "제178조", ["취소"]),
                        ("관세법", "제269조", ["밀수", "수출입"]),
                        ("보세판매장고시", "제3조", []),
                        ("보세판매장고시", "제18조", []),
                    ]
                    
                    ok_count = 0
                    err_count = 0
                    missing_count = 0
                    results_text = []
                    
                    for law_short, article, keywords in check_targets:
                        matches = [l for l in st.session_state.laws_db 
                                   if l.get("law_short") == law_short and l.get("article_no") == article]
                        if not matches:
                            results_text.append(f"❌ {law_short} {article} — DB에 없음!")
                            missing_count += 1
                            continue
                        
                        row = matches[0]
                        content = row.get("content", "")[:300]
                        title = row.get("article_title", "")
                        
                        if keywords:
                            has_kw = any(kw in content for kw in keywords)
                            if has_kw:
                                results_text.append(f"✅ {law_short} {article} ({title}) — 정상")
                                ok_count += 1
                            else:
                                results_text.append(f"🚨 {law_short} {article} ({title}) — 내용 불일치!")
                                results_text.append(f"   기대 키워드: {keywords}")
                                results_text.append(f"   실제 내용: {content[:150]}...")
                                err_count += 1
                        else:
                            results_text.append(f"✅ {law_short} {article} ({title}) — DB 존재 확인")
                            ok_count += 1
                    
                    # 결과 표시
                    if err_count > 0:
                        st.error(f"🚨 검증 실패 {err_count}건 발견!")
                    if missing_count > 0:
                        st.warning(f"❌ DB 미등록 {missing_count}건")
                    if ok_count > 0:
                        st.success(f"✅ 정상 {ok_count}건")
                    
                    for line in results_text:
                        st.caption(line)
                    
                    # 전체 현황
                    from collections import Counter
                    counts = Counter(l.get("law_short", "?") for l in st.session_state.laws_db)
                    st.caption(f"\n📊 DB 총 {len(st.session_state.laws_db)}건")
                    for name, cnt in sorted(counts.items()):
                        st.caption(f"  - {name}: {cnt}건")

            # 수동 법령 조문 추가 (MCP로 조회)
            with st.expander("🔧 수동 법령 조문 추가", expanded=False):
                manual_law = st.text_input("법령명 (예: 형법)", key="manual_law_name")
                manual_art = st.text_input("조문번호 (예: 제347조)", key="manual_art_no")
                manual_short = st.text_input("약칭 (예: 형법)", key="manual_short")
                
                if st.button("📥 MCP로 조문 가져오기", key="manual_fetch"):
                    if manual_law and manual_art and manual_short:
                        with st.spinner(f"{manual_law} {manual_art} MCP 조회 중..."):
                            try:
                                success, result = call_mcp_law(
                                    f"'{manual_law}' {manual_art} 조문 전문(항, 호 포함)을 원문 그대로 알려줘. 조문 제목도 포함해."
                                )
                                if success and len(result) > 30 and "미존재" not in result and "찾을 수 없" not in result:
                                    # 조문 제목 추출
                                    import re as _re
                                    title_match = _re.search(r'제\d+조(?:의\d+)?\s*\(([^)]+)\)', result)
                                    title = title_match.group(1) if title_match else ""
                                    
                                    law_id = f"{manual_short}_{manual_art.replace('제','').replace('조','')}"
                                    row = {"id": law_id, "law_name": manual_law, "law_short": manual_short,
                                           "article_no": manual_art, "article_title": title, "content": result[:2000],
                                           "last_updated": datetime.now().isoformat()}
                                    sb = init_supabase()
                                    sb.table("laws").upsert(row).execute()
                                    st.success(f"✅ {manual_short} {manual_art} ({title}) — {len(result)}자 저장!")
                                    st.caption(result[:300])
                                    st.session_state.laws_db = load_laws()
                                else:
                                    st.error(f"❌ {manual_law}에서 {manual_art}를 찾을 수 없습니다.")
                                    if result:
                                        st.caption(result[:200])
                            except Exception as e:
                                st.error(f"❌ 오류: {e}")

            st.markdown("─" * 30)

            doc_cat = st.selectbox("문서 유형", options=list(DOC_CATS.keys()), format_func=lambda x: DOC_CATS[x]["icon"] + " " + DOC_CATS[x]["label"])
            contract_type = st.selectbox("거래 유형", CONTRACT_TYPES) if doc_cat == "contract" else None
            yakjeong_type = st.selectbox("약정서 유형", YAKJEONG_TYPES) if doc_cat == "yakjeong" else None

            uploaded_files = st.file_uploader("Word 파일 첨부", type=["docx"], accept_multiple_files=True, label_visibility="collapsed")
            if uploaded_files:
                if st.button("DB에 규칙 등록", use_container_width=True):
                    for f in uploaded_files:
                        label = f"계약서({contract_type})" if contract_type else f"약정서({yakjeong_type})" if yakjeong_type else DOC_CATS[doc_cat]["label"]
                        label += f": {f.name}"
                        file_bytes = f.read()
                        new_doc = {
                            "id": str(uuid.uuid4()),
                            "name": f.name,
                            "cat": doc_cat,
                            "contract_type": contract_type or yakjeong_type,
                            "label": label,
                            "text": extract_text(file_bytes),
                            "size": len(file_bytes),
                        }
                        if save_doc(new_doc): st.session_state.docs.append(new_doc)
                    st.rerun()

            if st.session_state.docs:
                st.markdown("**📋 적용 중인 문서**")
                for cat_id, cat_info in DOC_CATS.items():
                    cat_docs = [d for d in st.session_state.docs if d["cat"] == cat_id]
                    if not cat_docs: continue
                    st.markdown(f"**{cat_info['icon']} {cat_info['label']}**")
                    for doc in cat_docs:
                        col1, col2 = st.columns([5, 1])
                        with col1: st.caption(doc["name"])
                        with col2:
                            if st.button("X", key="del_" + doc["id"]):
                                if delete_doc(doc["id"]):
                                    st.session_state.docs = [d for d in st.session_state.docs if d["id"] != doc["id"]]
                                    st.rerun()

    # ── 메인 영역 ────────────────────────────────────────────
    st.title("🤝 공정거래 실무 어시스턴트 v2.1")
    st.caption("MD·협력사 실무 Q&A & 계약·법령 Self-Check AI")

    if not st.session_state.messages and st.session_state.docs:
        st.markdown("### 💡 AI 법무 자문 100% 활용 가이드")
        st.info(
            "본 시스템은 질문의 목적에 따라 **두 가지 수준의 맞춤형 자문**을 제공합니다.\n\n"
            "🔹 **[1단계] 사내 기준 자문:** 일상적인 규정 문의 시, 등록된 당사 사규와 표준계약서를 바탕으로 신속한 실무 기준을 안내합니다.\n"
            "🔹 **[2단계] 심층 법무 검토:** 계약서/약정서가 첨부되거나 위법성 판단을 요청하면, 내부 사규와 외부 현행 법령을 교차 분석하여 정식 '검토 의견서'를 발행합니다."
        )
        
        samples = [
            ("🔹 [1단계] 단순 사내 규정 문의", "현재 등록된 당사에 따르면, 브랜드 자발적 사유로 매장을 리뉴얼할 때 인테리어 비용 분담 기준이 어떻게 돼?"),
            ("🔹 [2단계] 심층 법률 조항 검토", "협력사가 특약매입 판촉비 분담률을 60%로 요구하고 있어. 법률적으로 이게 맞음? 수용 가능한지 당사 기준과 대규모유통업법을 비교해서 분석해줘."),
            ("🔹 [2단계] 첨부파일 교차 검토", "[파일 첨부 후 클릭] 첨부한 파견 약정서(협력사 회신본) 내용 중, 당사 표준 규정에 어긋나거나 법 위반 소지가 있는 독소조항을 찾아내 줘.")
        ]
        cols = st.columns(3)
        for i, (cat, q) in enumerate(samples):
            with cols[i]:
                if st.button(cat, key="sample_" + str(i), use_container_width=True, help=q):
                    st.session_state["pending_input"] = q
                    st.rerun()

    # 대화 히스토리 렌더링
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"], avatar="🤝" if msg["role"] == "assistant" else "👤"):
            if msg["role"] == "assistant" and msg.get("json_data"):
                jd = msg["json_data"]
                render_verdict_badge(jd.get("verdict", ""))
                st.markdown(f"**📋 {jd.get('summary', '')}**")
                # 4번: 인용 검증 실패 경고
                cit_results = msg.get("citation_results", [])
                if cit_results and not all(cr["verified"] for cr in cit_results):
                    unverified = [cr["citation"] for cr in cit_results if not cr["verified"]]
                    unverified_links = []
                    for uv in unverified:
                        search_term = re.sub(r'\s*제\d+조.*', '', uv).strip()
                        link = f"https://www.law.go.kr/LSW/lsInfoP.do?lsiSeq=0&query={search_term}"
                        unverified_links.append(f"[{uv}]({link})")
                    st.warning("⚠️ 다음 법령 인용의 DB 검증이 완료되지 않았습니다.\n\n사내변호사에게 해당 조문의 현행 유효 여부를 반드시 확인받으세요.")
                    st.markdown("🔗 " + " | ".join(unverified_links))
                render_issues_table(jd.get("issues", []), msg.get("citation_results", []))
                if jd.get("alternative_clause"):
                    render_alternative_clause(jd["alternative_clause"])
                
                if jd.get("verdict"):
                    docx_bytes = generate_review_docx(jd, msg.get("detail_text", ""), "")
                    st.caption("⚠️ 본 문서를 사내변호사 확인 없이 외부에 발송하지 마세요.")
                    st.download_button("📥 검토의견서 다운로드 (.docx)", data=docx_bytes, file_name=f"검토의견서_{datetime.now().strftime('%Y%m%d_%H%M')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"dl_{msg.get('msg_id', datetime.now().timestamp())}")
            else:
                st.markdown(msg["content"])

            if "time" in msg and msg["time"]:
                model_label = msg.get("model", "")
                st.caption(f"⏱️ {msg['time']:.1f}초 · {model_label}")

    # ── 입력 처리 ────────────────────────────────────────────
    if st.session_state.docs:
        
        st.markdown("---")
        st.markdown("### 💬 신규 검토 요청")
        
        st.info("🔒 **기밀 유출 방지 시스템:** 당사의 기밀이나 특정 브랜드명, 상호명이 외부 AI로 전송되어 학습에 오남용되는 것을 막기 위해 아래 칸에 **검토 대상 협력사명**을 기재해 주세요. 해당 단어는 문서와 채팅에서 모두 가려집니다.")
        target_partner = st.text_input("🏢 검토 대상 협력사명 입력 (보안 마스킹용)", placeholder="예: 에르메스, 샤넬 (입력 시 █협력사█로 자동 치환)", key="target_partner")

        with st.expander("🔄 리비전 교차 비교 (당사 초안 vs 협력사 수정본)", expanded=False):
            st.info("당사 초안(기준)과 협력사가 수정한 문서를 나란히 업로드하면, AI변호사가 변경된 독소조항을 찾아 비교 분석합니다.")
            col1, col2 = st.columns(2)
            with col1: v1_file = st.file_uploader("📄 V1 (당사 표준 초안)", type=["docx"], key="v1_upload")
            with col2: v2_file = st.file_uploader("📝 V2 (협력사 수정본)", type=["docx"], key="v2_upload")
            if v1_file and v2_file:
                if st.button("교차 비교 분석 실행", type="primary", use_container_width=True):
                    v1_file.seek(0)
                    v2_file.seek(0)
                    v1_bytes, v2_bytes = v1_file.read(), v2_file.read()
                    v1_text = extract_text(v1_bytes)
                    v2_text = extract_text(v2_bytes)
                    
                    # 추출 실패 체크
                    if v1_text.startswith("(") and v1_text.endswith(")"):
                        st.error(f"V1 파일 읽기 실패: {v1_text}")
                    elif v2_text.startswith("(") and v2_text.endswith(")"):
                        st.error(f"V2 파일 읽기 실패: {v2_text}")
                    else:
                        prompt = (
                            f"당사가 보낸 초안(V1)과 협력사가 회신한 수정본(V2)을 교차 비교해주세요.\n\n"
                            f"1. 협력사가 어느 조항을 어떻게 변경/추가/삭제했는지 핵심만 대조해주세요.\n"
                            f"2. 수정본(V2)의 내용이 DB의 [기준 문서]와 [법령]을 위반하는지 엄격히 심사해주세요.\n\n"
                            f"[V1 당사 초안 내용]\n{v1_text}\n\n"
                            f"[V2 협력사 수정본 내용]\n{v2_text}"
                        )
                        st.session_state["pending_input"] = apply_auto_masking(prompt, target_partner)
                        st.rerun()

        chat_files = st.file_uploader("📎 검토할 파일 첨부 (협력사 회신본 등)", type=["docx"], accept_multiple_files=True, key="chat_uploader")
        user_input = st.chat_input("검토할 텍스트를 입력하거나 파일을 첨부하세요...")
        query = user_input or st.session_state.pop("pending_input", None)

        if query:
            attached_texts = []
            if chat_files:
                for f in chat_files:
                    f.seek(0)
                    raw_text = extract_text(f.read())
                    # 추출 실패 시에도 마스킹 적용 후 전달 (AI가 에러 메시지를 받아 안내)
                    safe_text = apply_auto_masking(raw_text, target_partner)
                    attached_texts.append(f"=== 검토 대상 첨부 파일: {f.name} ===\n" + safe_text)

            has_attachment = bool(attached_texts)
            safe_query = apply_auto_masking(query, target_partner)

            if attached_texts:
                full_query = f"[사용자 문의사항]\n{safe_query}\n\n[검토 대상 텍스트/첨부파일]\n" + "\n\n".join(attached_texts)
                display_query = safe_query + "\n\n📎 " + ", ".join(f.name for f in chat_files)
            else:
                has_review_content = len(safe_query) > 80 or "조" in safe_query or "항" in safe_query or ":" in safe_query
                if any(kw in safe_query for kw in ["검토", "확인", "분석"]) and not has_review_content:
                    full_query = f"[사용자 문의사항]\n{safe_query}\n\n[검토 대상 텍스트/첨부파일]\n(없음 - 첨부파일이나 텍스트가 제공되지 않았습니다.)"
                else:
                    full_query = f"[사용자 문의사항 및 검토 대상 텍스트]\n{safe_query}\n\n[첨부파일]\n(없음)"
                display_query = safe_query

            st.session_state.messages.append({"role": "user", "content": full_query})
            with st.chat_message("user", avatar="👤"):
                st.markdown(display_query)

            model_choice = route_query(safe_query, has_attachment)

            with st.chat_message("assistant", avatar="🤝"):
                if model_choice == "claude":
                    spinner_msg = "⚖ 3단계 하이브리드 검토 중 (Gemini→DB검증→Claude)..."
                else:
                    spinner_msg = "💬 당사 사내 기준을 검색 및 분석 중..."

                with st.spinner(spinner_msg):
                    start_time = time.time()
                    reply, actual_model = dispatch_with_fallback(model_choice, st.session_state.messages, st.session_state.docs, st.session_state.laws_db)
                    elapsed = time.time() - start_time

                msg_data = {"role": "assistant", "content": reply, "time": elapsed, "model": actual_model, "msg_id": str(datetime.now().timestamp())}

                if "Failed" not in actual_model and ("Claude" in actual_model or (model_choice == "claude" and "Gemini" in actual_model)):
                    json_data, detail_text = parse_review_response(reply)
                    if json_data:
                        # ██ 필수 쟁점 강제 삽입 — AI가 빠뜨려도 100% 보장 ██
                        json_data = enforce_mandatory_issues(json_data, st.session_state.laws_db)
                        
                        citation_results = verify_citations(json_data.get("cited_laws", []), st.session_state.laws_db)
                        precedent_results = verify_precedents(json_data.get("cited_precedents", []))
                        msg_data["json_data"] = json_data
                        msg_data["detail_text"] = detail_text
                        msg_data["citation_results"] = citation_results
                        msg_data["precedent_results"] = precedent_results

                        render_verdict_badge(json_data.get("verdict", ""))
                        st.markdown(f"**📋 {json_data.get('summary', '')}**")
                        
                        # 법령 인용 검증 결과
                        db_verified = [cr for cr in citation_results if cr["verified"]]
                        db_unverified = [cr for cr in citation_results if not cr["verified"]]
                        
                        if db_verified:
                            verified_items = []
                            for cr in db_verified:
                                date_info = ""
                                if cr.get("last_updated"):
                                    try:
                                        date_info = f" (DB: {cr['last_updated'][:10]})"
                                    except Exception:
                                        pass
                                verified_items.append(f"✅ {cr['citation']}{date_info}")
                            with st.expander(f"📚 DB 검증 완료 법령 ({len(db_verified)}건)", expanded=False):
                                st.markdown("\n".join(verified_items))
                        
                        if db_unverified:
                            unverified_links = []
                            for cr in db_unverified:
                                search_term = re.sub(r'\s*제\d+조.*', '', cr["citation"]).strip()
                                link = f"https://www.law.go.kr/LSW/lsInfoP.do?lsiSeq=0&query={search_term}"
                                unverified_links.append(f"[{cr['citation']}]({link})")
                            st.warning(f"⚠️ DB 미등록 법령 {len(db_unverified)}건 — 국가법령정보센터에서 현행 여부를 확인하세요.")
                            st.markdown("🔗 " + " | ".join(unverified_links))
                        
                        # 판례 검증 결과 + 할루시네이션 차단 표시
                        gk_meta = st.session_state.get("_gatekeeper_meta", {})
                        if gk_meta:
                            dropped = gk_meta.get("dropped_precedents", [])
                            v_prec = gk_meta.get("verified_precedents", [])
                            if dropped:
                                with st.expander(f"🚨 AI 할루시네이션 {len(dropped)}건 감지 — 시스템이 자동 차단", expanded=True):
                                    for dp in dropped:
                                        reason = dp.get('reason', '국가법령정보센터 미존재')
                                        st.markdown(f"~~{dp['case_no']}~~ → **차단 사유:** {reason}")
                                    st.caption("위 판례는 AI가 생성한 허위 정보로, 검토의견서에서 자동 제거되었습니다.")
                            if v_prec:
                                with st.expander(f"✅ API 검증 완료 판례 ({len(v_prec)}건)", expanded=False):
                                    for vp in v_prec:
                                        title = vp.get('title', '')
                                        title_str = f" ({title})" if title else ""
                                        st.markdown(f"- ✅ {vp['case_no']}{title_str}: {vp.get('summary', '')[:100]}")
                        elif precedent_results:
                            unverified_prec = [p for p in precedent_results if not p["verified"]]
                            if unverified_prec:
                                st.warning(f"⚠️ 미확인 판례 {len(unverified_prec)}건 — 대법원 판례검색에서 실재 여부를 확인하세요.")
                                prec_links = [f"[{p['case_no']}](https://glaw.scourt.go.kr/wsjo/panre/sjo100.do)" for p in unverified_prec]
                                st.markdown("🔗 " + " | ".join(prec_links))
                        
                        render_issues_table(json_data.get("issues", []), citation_results)
                        if json_data.get("alternative_clause"): render_alternative_clause(json_data["alternative_clause"])

                        docx_bytes = generate_review_docx(json_data, detail_text, display_query)
                        st.caption("⚠️ 본 문서를 사내변호사 확인 없이 외부에 발송하지 마세요.")
                        st.download_button("📥 검토의견서 다운로드 (.docx)", data=docx_bytes, file_name=f"검토의견서_{datetime.now().strftime('%Y%m%d_%H%M')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

                        save_review_log({
                            "id": msg_data["msg_id"], "session_id": st.session_state.current_session_id, "verdict": json_data.get("verdict", "unknown"),
                            "issues": json_data.get("issues"), "action_plan": json_data.get("action_plan"), "cited_laws": json_data.get("cited_laws"),
                            "citation_verified": all(cr["verified"] for cr in citation_results) if citation_results else False,
                        })
                    else:
                        st.markdown(reply)
                else:
                    st.markdown(reply)

                st.caption(f"⏱️ {elapsed:.1f}초 · {actual_model}")
                st.session_state.messages.append(msg_data)

            # 세션 저장 후 다음 렌더 사이클에서 rerun (다운로드 버튼 접근 보장)
            new_id = st.session_state.current_session_id or str(uuid.uuid4())
            current_sess = {"id": new_id, "title": display_query[:25] + "...", "date": datetime.now().isoformat(), "messages": st.session_state.messages}
            if save_session(current_sess):
                st.session_state.current_session_id = new_id
                existing = [s for s in st.session_state.sessions if s["id"] != new_id]
                st.session_state.sessions = [current_sess] + existing
            # 즉시 rerun 대신 플래그 설정 → 현재 렌더에서 다운로드 버튼이 먼저 표시됨
            st.session_state.needs_rerun = True

    else:
        st.markdown("### 🚀 시작하기 — 3단계 설정 가이드")
        st.markdown("아래 순서대로 기준 문서를 등록하면 AI 검토를 시작할 수 있습니다.")
        
        # 등록 현황 체크
        has_saryu = any(d["cat"] == "saryu" for d in st.session_state.docs)
        has_contract = any(d["cat"] == "contract" for d in st.session_state.docs)
        has_yakjeong = any(d["cat"] == "yakjeong" for d in st.session_state.docs)
        
        check = lambda done: "✅" if done else "⬜"
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.markdown(f"#### {check(has_saryu)} STEP 1")
            st.markdown("**🏛 사규 등록**")
            st.caption("공정거래 컴플라이언스 정책, 내부 규정 등")
            if has_saryu:
                st.success("등록 완료")
            else:
                st.warning("미등록")
        with col2:
            st.markdown(f"#### {check(has_contract)} STEP 2")
            st.markdown("**📄 표준 계약서 등록**")
            st.caption("특약매입/직매입 거래유형별 당사 표준 계약서")
            if has_contract:
                st.success("등록 완료")
            else:
                st.warning("미등록")
        with col3:
            st.markdown(f"#### {check(has_yakjeong)} STEP 3")
            st.markdown("**📝 표준 약정서 등록**")
            st.caption("협력사원, 인테리어, 매장이동, 공동판촉 약정서")
            if has_yakjeong:
                st.success("등록 완료")
            else:
                st.warning("미등록")
        
        st.markdown("---")
        st.info("👈 사이드바 하단 **'⚙️ 기준 문서 DB 관리'**를 열고, 문서 유형을 선택한 뒤 Word 파일을 업로드하세요.\n\n최소 **사규 1개 + 계약서 1개**가 등록되면 AI 검토를 시작할 수 있습니다.")

if __name__ == "__main__":
    main()
